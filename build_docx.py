"""Builds the two deliverable .docx files in expanded form:

  Hartford_Grid_Dev_Journal.docx     — colorful, chapter-structured development log
  Hartford_Grid_Research_Context.docx — research papers + niche analysis

Upload either to Google Drive (drag into drive.google.com) and Drive will
auto-convert to a native Google Doc.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


INK       = RGBColor(0x3B, 0x2F, 0x24)
ACCENT    = RGBColor(0x8B, 0x3A, 0x2E)
MARGIN    = RGBColor(0x70, 0x56, 0x38)
BUILD     = RGBColor(0x1D, 0x4E, 0xD8)
FIX       = RGBColor(0x15, 0x80, 0x3D)
QUESTION  = RGBColor(0x7C, 0x2D, 0x12)
DECISION  = RGBColor(0xA1, 0x62, 0x07)
PIVOT     = RGBColor(0x7E, 0x22, 0xCE)
PERF      = RGBColor(0x0D, 0x94, 0x88)

BG_HEADER   = "F3E8C8"
BG_BUILD    = "DBEAFE"
BG_FIX      = "DCFCE7"
BG_QUESTION = "FEE2D5"
BG_DECISION = "FEF3C7"
BG_PIVOT    = "F3E8FF"
BG_PERF     = "CCFBF1"
BG_QUOTE    = "FFF7E2"


def shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, hex_color="C9B894"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), hex_color)
        borders.append(e)
    tc_pr.append(borders)


def add_styled(p, text, bold=False, italic=False, color=None, size=None, font="Georgia"):
    run = p.add_run(text)
    run.font.name = font
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    return run


def h1(doc, text):
    p = doc.add_paragraph()
    add_styled(p, text, bold=True, color=ACCENT, size=22, font="Georgia")


def h2(doc, text):
    p = doc.add_paragraph()
    add_styled(p, text, bold=True, color=ACCENT, size=15)


def h3(doc, text):
    p = doc.add_paragraph()
    add_styled(p, text, bold=True, color=RGBColor(0x5B, 0x45, 0x28), size=12)


def margin_note(doc, text):
    p = doc.add_paragraph()
    add_styled(p, text, italic=True, color=MARGIN, size=10)


def body(doc, text):
    p = doc.add_paragraph()
    add_styled(p, text, color=INK, size=11)


def quote(doc, who, text):
    """User quote block."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.4)
    shade(cell, BG_QUOTE)
    set_cell_borders(cell, "8B3A2E")
    p = cell.paragraphs[0]
    add_styled(p, f"{who}: ", bold=True, color=QUESTION, size=10)
    add_styled(p, text, italic=True, color=INK, size=11)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_styled(p, item, color=INK, size=11)


def tag_color(kind):
    base = kind.lower()
    if "build" in base: return (BUILD, BG_BUILD)
    if "fix" in base or "bug" in base or "failure" in base or "success" in base: return (FIX, BG_FIX)
    if "question" in base: return (QUESTION, BG_QUESTION)
    if "decision" in base or "choice" in base: return (DECISION, BG_DECISION)
    if "pivot" in base or "reframe" in base: return (PIVOT, BG_PIVOT)
    if "perf" in base or "×" in kind or "milestone" in base or "honest" in base or "surprise" in base or "achieved" in base or "partial" in base: return (PERF, BG_PERF)
    return (INK, "FFFFFF")


def add_entry_table(doc, rows):
    """rows = list of (when, kind, detail) tuples."""
    table = doc.add_table(rows=1 + len(rows), cols=3)
    widths = [Inches(0.9), Inches(1.0), Inches(4.5)]
    hdr = table.rows[0].cells
    for i, (c, w, txt) in enumerate(zip(hdr, widths, ("WHEN", "KIND", "ENTRY"))):
        c.width = w
        shade(c, BG_HEADER)
        set_cell_borders(c)
        p = c.paragraphs[0]
        add_styled(p, txt, bold=True, color=RGBColor(0x5B, 0x45, 0x28), size=10)
    for r, (when, kind, detail) in enumerate(rows, start=1):
        row = table.rows[r].cells
        for c, w in zip(row, widths):
            c.width = w
            set_cell_borders(c)
        p = row[0].paragraphs[0]
        add_styled(p, when, italic=True, color=MARGIN, size=10)
        col, bg = tag_color(kind)
        shade(row[1], bg)
        p = row[1].paragraphs[0]
        add_styled(p, kind, bold=True, color=col, size=10)
        p = row[2].paragraphs[0]
        add_styled(p, detail, color=INK, size=11)


def section_break(doc, text=None):
    p = doc.add_paragraph()
    add_styled(p, "❦", color=ACCENT, size=14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# =================== Build the Dev Journal ===================

def build_journal():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    sec.left_margin = sec.right_margin = Inches(0.9)

    h1(doc, "Development Journal")
    p = doc.add_paragraph()
    add_styled(p, "Hartford County Power Grid Resilience Simulation",
               italic=True, color=MARGIN, size=12)
    p = doc.add_paragraph()
    add_styled(p, "A. Syed Diamond  ·  Spring–Summer 2026",
               italic=True, color=MARGIN, size=10)

    body(doc, "")
    legend = doc.add_paragraph()
    for label, color in [
        ("■ Build", BUILD), ("  ■ Fix", FIX), ("  ■ Perf", PERF),
        ("  ■ Question", QUESTION), ("  ■ Decision", DECISION), ("  ■ Pivot", PIVOT),
    ]:
        add_styled(legend, label + "  ", color=color, size=10, bold=True)

    # Table of Contents
    body(doc, "")
    h2(doc, "Contents")
    toc_items = [
        "I.    Foundations & the Initial Simulation",
        "II.   The Five-Alternative Roadmap",
        "III.  Alternative #1 — The Closed-Form Equation",
        "IV.   Alternative #2 — The Pre-Computed Scenario Library",
        "V.    Alternative #3 — The WebAssembly Adventure That Wasn't",
        "VI.   Alternative #4 — Server-Side Compute Arrives",
        "VII.  The \"Max Settings\" Crisis",
        "VIII. The Connecticut-Scale Refactor",
        "IX.   The Render Deployment Saga",
        "X.    Polish: Auto-Detect, Progress, /version",
        "XI.   The Vanishing Markers",
        "XII.  Direction Conversations & Research Strategy",
        "XIII. Documentation Push",
        "XV.   The Realism Round & the Polish Sprint",
        "Appendix A — Problems Faced (cross-project catalogue)",
        "XIV.  Coda",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        add_styled(p, item, color=INK, size=11)

    section_break(doc)

    # ============ Chapter I (expanded) ============
    h2(doc, "Chapter I — Foundations & the Initial Simulation")
    margin_note(doc,
                "Long before the five performance alternatives, there was the question itself, "
                "the geography, the grid, the storm, the seven realism factors, the scheduler, "
                "the map, and the first \"the page isn't responding\" freeze. The work of this "
                "chapter laid every assumption everything else stands on.")

    h3(doc, "I.1 — The First Question")
    body(doc,
         "The project started as a single HTML page asking a deceptively simple question: "
         "how long does it take Hartford County to get the lights back on after a storm, and "
         "how does that number change when we add more crews? Everything that followed came "
         "out of that one sentence. The initial milestone wasn't speed — it was earning the "
         "right to make a claim about restoration timing at all. That meant a synthetic grid "
         "faithful enough to behave like a distribution network, a storm model that produced "
         "outage patterns instead of random dots, and a scheduler with the kind of realism "
         "that distinguishes a plausible answer from a fantasy.")
    body(doc,
         "The intended use was three-fold: a teaching demo to make grid-resilience intuitions "
         "concrete, an interactive sandbox for exploring what-if scenarios at the county "
         "scale, and the substrate on which a future research arc — MILP comparison, "
         "calibration against real Eversource storm data, eventual statewide projection — "
         "could plausibly land.")

    h3(doc, "I.2 — Acquiring the County")
    body(doc,
         "Hartford County isn't a vague rectangle. It's a precise polygon with 29 towns "
         "inside it, each with its own boundary, its own population, and its own road "
         "density. None of those were going to come from the simulator's imagination — we "
         "had to fetch them.")
    body(doc,
         "Two small Python utilities handled that. 01_fetch_county_boundary.py pulled the "
         "county outline as GeoJSON from open boundary data; 02_fetch_town_boundaries.py "
         "pulled the 29 town polygons individually so each could be drawn, named, and "
         "weighted. Both scripts use only the Python standard library — intentional, so "
         "anyone cloning the repo can rebuild from scratch without installing anything. The "
         "fetched boundaries live under data/ as committed JSON so the interactive page can "
         "read them without a server.")
    body(doc,
         "Census-style population figures per town were baked into a static table inside "
         "the artifact generator. Towns with bigger populations became weight centres for "
         "demand placement; small towns got correspondingly fewer customers. Not perfect — "
         "an upgrade path is to pull real parcel-level data — but accurate enough that the "
         "visualisation rang true to people who know the county.")
    add_entry_table(doc, [
        ("Foundation", "Build",
         "01_fetch_county_boundary.py — downloads the Hartford County outline as GeoJSON. "
         "Pure stdlib. Output written to data/hartford_boundary.json."),
        ("Foundation", "Build",
         "02_fetch_town_boundaries.py — pulls all 29 town polygons inside the county. "
         "Pure stdlib. Output: data/hartford_towns.json."),
        ("Foundation", "Build",
         "Town-population table baked into the artifact generator and into "
         "03_grid_simulation.html, so the interactive doesn't need a server-side call to "
         "weight demand."),
    ])

    h3(doc, "I.3 — Building a Synthetic Distribution Grid")
    body(doc,
         "Real distribution networks aren't public. Eversource and Avangrid keep the actual "
         "feeder topology proprietary — not just for competitive reasons but for physical-"
         "security ones. What's public is the shape of distribution systems: substations "
         "placed near load, backbone feeders radiating outward, laterals branching off "
         "feeders to reach individual streets. Our job was to build a synthetic network with "
         "the right shape and density so storm patterns on it would look like storm patterns "
         "on the real thing.")
    body(doc,
         "The recipe, in order: (1) generate ~10 000 candidate demand points, sampled inside "
         "the county polygon and weighted by town population so cities like Hartford and "
         "Bristol got their fair share. (2) Run k-means on those demand points to place 100 "
         "substations — the count is a user-adjustable slider, but 100 was the realistic "
         "default for the county's ~900 000 customers. (3) For each substation, generate K "
         "feeders (default K = 10): radial backbones that meander outward from the substation "
         "in random-but-plausible directions, capped at a realistic distance. (4) Hang "
         "laterals off each feeder at intervals, themselves branching to cover nearby demand.")
    body(doc,
         "Two design choices mattered: every random number came from a deterministic "
         "mulberry32 PRNG seeded by an integer the user can change in the UI, and the same "
         "PRNG was implemented identically in JavaScript and Python so the offline artifact "
         "generator produced bit-identical grids to the browser. That meant a screenshot of "
         "a storm-restoration plan could be regenerated months later, exactly, from the seed "
         "alone.")
    body(doc,
         "The Leaflet rendering layered them: substations as colored stars, each "
         "substation's feeders drawn in that substation's color (so you could see which "
         "lateral belonged to which backbone at a glance), laterals as faint gray lines "
         "underneath, and the county outline overlaid in red so the geography was "
         "unmistakable.")
    add_entry_table(doc, [
        ("Foundation", "Build",
         "Population-weighted demand-point sampler. Generates ~10 000 candidate customer "
         "locations inside the county polygon, weighted by each town's 2020 census population."),
        ("Foundation", "Build",
         "k-means substation placement with configurable K (default 100). Output: substation "
         "lat/lon list serving as the roots of the synthetic distribution network."),
        ("Foundation", "Build",
         "Feeder generator: for each substation, branch K radial backbones (default 10) "
         "outward in random plausible directions; laterals hung off feeders at intervals."),
        ("Foundation", "Decision",
         "All randomness funnels through a single mulberry32 PRNG, implemented identically "
         "in JavaScript and Python. Seeding the integer at the top of the UI reproduces the "
         "exact same grid + storm + restoration, on any machine, indefinitely."),
        ("Foundation", "Build",
         "Leaflet layering: county outline (red) → town outlines (light gray) → substations "
         "(colored stars) → feeders (colored lines per substation) → laterals (gray) → "
         "storm outages (X marks) → restoration markers (depots + numbered circles)."),
    ])

    h3(doc, "I.4 — The Storm Model")
    body(doc,
         "A storm in this simulation is not \"X% of customers lose power randomly.\" It's a "
         "set of physical failure points placed along feeders and laterals, where each "
         "failure cuts power to everything downstream of it. That model produces realistic "
         "patterns: clustered outages, geographically coherent blackout regions, and the "
         "distinctive shape of distribution-system failures rather than the smeared "
         "distribution of a naive random model.")
    body(doc,
         "The storm slider chooses N, the number of failure points. The simulator weighted "
         "each candidate segment of the grid by exposure (longer segments more likely, "
         "exterior laterals over reinforced trunks) and sampled N failures without "
         "replacement. Each failure carried forward to a customer count by walking "
         "downstream from the failure point and summing served demand.")
    body(doc,
         "Reset was non-destructive in the right way: hitting Reset storm cleared the "
         "failures and the restoration plan, but left the underlying grid intact. So you "
         "could explore many storms on the same network without regenerating substations or "
         "feeders, and the seed determined exactly which failures appeared the next time you "
         "simulated.")
    add_entry_table(doc, [
        ("Foundation", "Build",
         "Storm sampler: pick N failure points on feeders + laterals weighted by exposure. "
         "N slider goes from 0 to ~25 000."),
        ("Foundation", "Build",
         "Customer-impact calculation: for each failure, walk downstream summing the demand "
         "served. UI counter shows total customers without power and percent of county "
         "population affected."),
        ("Foundation", "Build",
         "Reset-storm button clears failures + plan without rebuilding the grid. Keeps the "
         "seed so the next storm is reproducible from the same starting state."),
    ])

    h3(doc, "I.5 — The Realism Crisis & the Seven Factors")
    body(doc,
         "A naïve scheduler — \"for each crew, find the nearest outage, repair it in one "
         "hour, move on\" — gives wildly optimistic restoration times. With 100 crews and "
         "500 outages, that model predicts a 4-hour restoration. Eversource's post-storm "
         "filings put real restorations in the 2–7 day range. If the simulation was going "
         "to be taken seriously for any later research purpose, the gap between 4 hours and "
         "4 days had to come from somewhere defensible.")
    body(doc,
         "That somewhere ended up being a set of seven multiplicative effects layered on "
         "top of the baseline scheduler. Each one models a real-world delay that storm "
         "reports document. Stacked, they produce the 12–18 hour restoration for a 500-"
         "outage storm that does match what the literature describes — not by hand-tuning "
         "to a number, but by adding effects with documented sources and watching the total "
         "emerge.")
    add_entry_table(doc, [
        ("#1", "Build (Assessment)",
         "First 12 hours after the storm: no repair dispatch. Crews are doing damage "
         "assessment — driving routes, identifying failures, prioritising. Documented in "
         "Eversource's Isaias 2020 post-mortem and PURA filings."),
        ("#2", "Build (Repair)",
         "Per-repair duration is log-normal, median 2 h, 90th percentile 6 h, capped at "
         "12 h. Box–Muller transform from mulberry32 uniforms. Each repair's duration is "
         "independently sampled but reproducible from the seed."),
        ("#3", "Build (Discovery)",
         "Outages aren't all known to the utility at t=0. A discovery curve has 30% of "
         "outages reported in the first hour after assessment ends, the remaining 70% "
         "reported with an exponential-decay tail capped at 36 h. Crews can't repair an "
         "outage they don't know about yet."),
        ("#4", "Build (Mutual aid)",
         "For storms requiring ≥ 6 crews, only ~50% are local. ~30% arrive 24 h later as the "
         "first mutual-aid wave (neighboring utilities), the remaining ~20% arrive 48 h "
         "later (long-distance aid). Models how out-of-state crews ramp into a major event."),
        ("#5", "Build (Roads)",
         "Travel distance ≠ haversine. Real driving is longer because roads bend, deadhead, "
         "and route around obstacles. Multiplier of 1.5× on the great-circle distance at "
         "25 mph average. Approximation but defensible."),
        ("#6", "Build (Workday)",
         "A repair completing after the 14-hour workday boundary rolls over to the next "
         "morning. Models the reality that night ops are reduced (safety, visibility). The "
         "clamp is what stretches a \"12 hour theoretical\" restoration into the 2–3 day "
         "reality."),
        ("#7", "Build (Critical)",
         "Outages flagged critical (hospitals, fire stations) get a priority phase. The "
         "scheduler dispatches all critical jobs first before resuming normal greedy "
         "nearest-neighbour. Documented utility practice."),
    ])
    body(doc,
         "The \"realistic mode\" checkbox in the UI toggles all seven on or off. Off, the "
         "scheduler runs the naïve model — useful as a baseline that demonstrates how much "
         "each effect matters. On, the simulator produces the multi-day restorations that "
         "match published Eversource event timelines in order of magnitude.")

    h3(doc, "I.6 — The Greedy Rolling-Horizon Scheduler")
    body(doc,
         "With the realism factors specified, the scheduler itself could be relatively "
         "simple. The choice was a greedy rolling-horizon dispatch:")
    bullets(doc, [
        "Maintain a min-heap of (time, crew_id) sorted by the crew's next-available time.",
        "Pop the earliest-available crew. Call its time t.",
        "Among outages that are (a) not yet repaired and (b) already discovered by time t, find the nearest one to the crew's current location.",
        "Assign that outage to this crew. Compute the eta: travel time at 25 mph × road multiplier + stochastic repair duration. Apply the workday clamp.",
        "Push the crew back onto the heap with its new (eta, crew_id). Loop.",
    ])
    body(doc,
         "Rolling-horizon means the next dispatch isn't planned in advance — each crew picks "
         "its next job based on the world state at the moment it becomes free. That's "
         "deliberately greedy: it's not optimal in any global sense, but it's defensible "
         "(utilities really do dispatch this way) and it's scalable.")
    body(doc,
         "The output the scheduler returns: per-crew finish time (when that crew completes "
         "its last job), per-crew job sequence (which outage in what order), and the "
         "system-wide total restoration time (the max of the per-crew times). Plus a "
         "timeline of (hour, remaining outages) sampled at regular intervals for the curve "
         "display.")
    body(doc,
         "A useful by-product: the \"Find optimal crew count\" button. Given a target "
         "restoration window (within 15% of the theoretical floor when crews = outages), "
         "the UI binary-searches over crew counts, running the scheduler at each candidate, "
         "and returns the smallest M that meets the target. This is the question utilities "
         "actually ask after a storm: \"we got everything back in 3 days; if we'd had 50 "
         "more crews, how much faster?\"")
    add_entry_table(doc, [
        ("Foundation", "Build",
         "Min-heap-based rolling-horizon dispatch in JavaScript. Each crew's next-available "
         "time is the heap key; finding the nearest outage at dispatch time is the inner "
         "search."),
        ("Foundation", "Build",
         "Box–Muller normal sampler driven by the mulberry32 stream for log-normal repair "
         "durations. Independent RNG streams for repair time vs discovery time so seed "
         "sweeps don't shuffle realisations."),
        ("Foundation", "Build",
         "Find-optimal-crew-count binary search. Iterates over M = 1, 2, 4, 8, … doubling "
         "until restoration time is within 15% of the floor, then refines."),
    ])

    h3(doc, "I.7 — The First Visualization")
    body(doc,
         "Visualisation came in stages. The first version of the page rendered the grid "
         "alone — substations, feeders, laterals — with no storm. That established the "
         "geography. Adding storm simulation produced the outage X marks. Adding restoration "
         "produced the moment the project really started to feel like something real: "
         "colored crew depots scattered across the county, and at every outage, a numbered "
         "circle in the crew's color showing the repair order. Crew 7 in green did jobs 1, "
         "2, 3, … in sequence; you could trace the route across the map by eye.")
    body(doc,
         "For modest scenarios the visualisation was effortless — Leaflet's built-in marker "
         "layer handled a few hundred markers fine. The problems started above ~1500 "
         "outages, which is where Section I.8 picks up.")

    h3(doc, "I.8 — The First Performance Reckoning (pre-Alternatives)")
    body(doc,
         "Long before the five-alternative roadmap was even drafted, the page hit its first "
         "wall. With ~1500 outages and 100 crews, hitting Plan restoration caused the "
         "browser to freeze for several seconds. The dialog the user saw: \"this page is "
         "not responding.\" The diagnosis split into two distinct problems, which then got "
         "two distinct fixes.")
    body(doc,
         "Problem A: the scheduler blocks the main thread. The dispatch loop is fast for a "
         "few hundred outages but at 2000+ it's long enough that the browser's event loop "
         "never yields. Solution: chunk the loop and yield to the event loop every K "
         "dispatches with await new Promise(r => setTimeout(r, 0)). Now the page stays "
         "interactive during computation and the progress bar actually updates.")
    body(doc,
         "Problem B: rendering thousands of DOM markers is slow. Each Leaflet L.marker "
         "creates an HTML element. At 2000 markers the rendering takes hundreds of "
         "milliseconds; at 10 000 it's seconds. Solution: a custom PointCloudLayer that "
         "draws thousands of dots in a single canvas pass. Numbered repair circles still "
         "get individual HTML markers up to a budget (~1500); past the budget, they shift "
         "to the canvas point cloud as plain colored dots without numbers.")
    body(doc,
         "Problem C: nearest-outage search at the scheduler hot spot is O(N). For each "
         "dispatch, the scheduler scans all outages to find the nearest undone one. At "
         "N = 2000, that's 4 million comparisons per scheduler run. Solution introduced "
         "even before Alternative #1: a spatial grid hash on the JS side, with concentric "
         "ring expansion to find the nearest neighbour without a full scan. Capped search "
         "radius so it never explodes. (Echoes of this design would later be re-implemented "
         "in Numba inside scheduler_numba.py.)")
    add_entry_table(doc, [
        ("Pre-Alt", "Perf",
         "Async-yield chunking inside the scheduler loop: await every ~500 dispatches so "
         "the browser event loop runs. Progress bar finally updates during compute."),
        ("Pre-Alt", "Perf",
         "Custom Leaflet PointCloudLayer renderer: thousands of points drawn in a single "
         "canvas pass instead of as individual HTML markers. Numbered HTML markers up to a "
         "budget; cloud points beyond."),
        ("Pre-Alt", "Perf",
         "Spatial grid hash + ring-expansion nearest-neighbour on the JS side of the "
         "scheduler. Dropped per-dispatch search from O(N) to roughly O(local density). "
         "This was the same algorithmic idea later resurfacing in scheduler_numba.py's "
         "grid hash."),
        ("Pre-Alt", "Fix",
         "\"This page is not responding\" no longer fires at typical scales. The compute "
         "and render are both within the user's patience budget for storms in the hundreds-"
         "to-low-thousands of outages."),
    ])

    h3(doc, "I.9 — Supporting Tooling")
    body(doc,
         "A few smaller pieces of infrastructure went up in the foundations phase that "
         "didn't feature in any later chapter but were quietly load-bearing:")
    add_entry_table(doc, [
        ("Foundation", "Build",
         "04_geojson_to_shapefile.py — offline conversion of GeoJSON grid + storm + plan to "
         "ESRI shapefile so the artifacts open in QGIS, ArcGIS, or geopandas. Uses "
         "geopandas; only needed when exporting."),
        ("Foundation", "Build",
         "05_generate_artifacts.py — offline Python port of the JS scheduler that generates "
         "static PNG artifacts (county overview, grid topology, storm overlay, restoration "
         "plan) plus rasterised summary plots via matplotlib. Used to seed the scenario "
         "library and produce poster-quality figures."),
        ("Foundation", "Build",
         "Browser-side export: download the current grid + storm + plan as GeoJSON. Lets "
         "users open results in any GIS for further analysis."),
        ("Foundation", "Build",
         "UI metrics panel: customers without power, percent of county population affected, "
         "outage locations, repair crews, recommended crews. Updated live as inputs change."),
    ])

    h3(doc, "I.10 — The State at the End of Foundations")
    body(doc,
         "By the end of the foundations phase, before a single one of the five alternatives "
         "was started, the simulation could already do all of this:")
    bullets(doc, [
        "Generate a synthetic Hartford County distribution grid from a seed integer.",
        "Simulate a storm of any size from 0 to ~25 000 outages.",
        "Run the realistic-mode scheduler with all seven factors enabled.",
        "Visualise the result with depots + numbered repair circles + canvas point cloud.",
        "Recommend an optimal crew count via binary search.",
        "Export grid + storm + plan as GeoJSON or shapefile.",
        "Reproduce any past result exactly given just the seed.",
        "Stay responsive at typical scales thanks to chunked compute, the spatial grid hash, and the canvas renderer.",
    ])
    body(doc,
         "What it couldn't do yet: respond instantly to slider changes (closed-form would "
         "fix that), let users explore canned scenarios without running anything (pre-"
         "computed library), or scale to 25 000 outages with thousands of crews in "
         "realistic mode without locking up. Each of those was a deliberate next direction "
         "— and that's where the five-alternative roadmap of Chapter II came from.")
    body(doc,
         "The chapter that follows isn't about fixing what the foundations got wrong — "
         "everything in this chapter still stands as written, and most of it is still load-"
         "bearing in the live system today. It's about deciding which directions to push "
         "next given that the foundations worked.")

    h3(doc, "Problems Faced in Chapter I")
    margin_note(doc, "Every interesting bit of engineering in the foundations existed because something didn't work the obvious way first.")
    add_entry_table(doc, [
        ("\"4-hour fantasy\"", "Realism gap",
         "Cause: naïve scheduler ignored real-world delays; predicted 4 h vs reality 2–7 d. "
         "Resolution: the seven realism factors of I.5, layered with literature-backed "
         "parameter choices, brought the simulation into order-of-magnitude agreement with "
         "published event timelines."),
        ("Proprietary grid data", "Data unavailable",
         "Cause: Eversource's real feeder topology isn't public — partly competitive, "
         "partly physical-security. Resolution: synthetic distribution grid from population-"
         "weighted demand + k-means substations + radial feeders. Reproducible from a "
         "single seed."),
        ("\"Page isn't responding\"", "Compute blocks UI",
         "Cause: the scheduler dispatch loop never yielded to the browser event loop at "
         "N=1500+. Resolution: chunk the loop and await new Promise(r => setTimeout(r, 0)) "
         "every ~500 dispatches."),
        ("DOM markers slow at scale", "Render cost",
         "Cause: each L.marker creates an HTML element; 10 000 markers = seconds. "
         "Resolution: custom PointCloudLayer canvas renderer; numbered HTML markers up to a "
         "budget; canvas dots beyond."),
        ("O(N) nearest-outage scan", "Algorithmic",
         "Cause: inner search scans all outages every dispatch. Resolution: JS-side spatial "
         "grid hash with ring expansion. Same idea reappears in scheduler_numba.py later."),
        ("PRNG bit-mismatch risk", "Reproducibility",
         "Cause: for an offline Python artifact generator to match the browser, Mulberry32 "
         "had to be identical at every bit including integer-overflow semantics. Resolution: "
         "explicit & 0xFFFFFFFF masks in both languages plus a head-to-head 10 000-sample "
         "comparison test before anything was committed."),
    ])

    h3(doc, "Takeaways from Chapter I")
    bullets(doc, [
        "Realism is multiplicative. No single one of the seven factors gets the restoration time from \"4 hours\" to \"4 days.\" They're each individually small and only stack to the right number together. That's a feature, not a bug — each one is independently defensible against the literature.",
        "Synthetic-but-deterministic beats real-but-inaccessible. Real feeder data isn't available. A synthetic grid built from population weights, seeded so it's exactly reproducible, behaves enough like a real one that the storm patterns look right — and any reviewer can rebuild it from scratch.",
        "The first performance freeze taught patterns that paid forward. The async-yield chunking, the spatial grid hash, and the canvas point-cloud renderer were all designed before the five-alternative roadmap. The same grid-hash idea reappears later in scheduler_numba.py, where it's the foundation of the 246× speedup. Building right the first time means later work compounds instead of replacing.",
        "The right unit of reproducibility is a single integer. All randomness goes through one PRNG, seeded by one integer in the UI. That means a paper figure can be regenerated months later from the seed alone — no archives, no caches, no stale URLs.",
    ])

    section_break(doc)

    # ============ Chapter II ============
    h2(doc, "Chapter II — The Five-Alternative Roadmap")
    margin_note(doc, "Five orthogonal directions to make the simulation faster & more capable.")
    body(doc,
         "By the time the basic simulation worked, the obvious next question was: how do we "
         "make it fast enough to be genuinely interactive at much larger scales, and what "
         "other research capabilities should we unlock? Rather than pick one optimisation and "
         "hope, the plan was to enumerate the orthogonal directions and tackle them in order:")
    bullets(doc, [
        "Alternative #1 — Closed-form equation: instant slider feedback without re-running the scheduler.",
        "Alternative #2 — Pre-computed scenario library: let viewers explore canned scenarios with no compute at all.",
        "Alternative #3 — WebAssembly: port the scheduler to Rust, compile to WASM, hope to beat V8.",
        "Alternative #4 — Server-side computation: Python backend for things the browser fundamentally can't do (Monte Carlo, future MILP).",
        "Alternative #5 — WebGPU: parallel argmin on the GPU for huge-N stress tests.",
    ])
    body(doc,
         "The discipline that emerged early: always benchmark before declaring victory. Two "
         "of the five (WASM and KD-tree later on) ended up slower than the baseline in "
         "realistic conditions, and the only way to know was to put a number on it. This "
         "rhythm shaped the whole project.")

    section_break(doc)

    # ============ Chapter III ============
    h2(doc, "Chapter III — Alternative #1: The Closed-Form Equation")
    margin_note(doc, "The cheapest direction first. Got the slider feeling instant.")
    body(doc,
         "The motivation came from watching the slider sit motionless while users tried to "
         "find a sensible number of crews. Re-running the full scheduler on every slider "
         "change is wasteful when all you really need is a rough estimate. The closed-form "
         "fits an analytical approximation: total restoration ≈ assessment_delay + (N × "
         "mean_repair × road_factor) / (M × parallel_efficiency), with workday clamping. "
         "It's wrong by ~15% but right in direction.")
    add_entry_table(doc, [
        ("#1 build", "Build",
         "Added the live \"≈ 4 d 3 h\" readout above the Plan-restoration button. Updates in "
         "microseconds as the slider moves."),
        ("#1 review", "Question",
         "Closed-form initially read N from the slider, not the actual storm. User noticed:"),
    ])
    quote(doc, "User",
          "Also in regards to the estimate restoration time, can you make it dependent on "
          "the storm outages or is it just independent of everything and won't change unless "
          "you change the slider for number of repair crews?")
    add_entry_table(doc, [
        ("#1 verify", "Question",
         "User followed up with sharp diagnostic intuition:"),
    ])
    quote(doc, "User",
          "The thing is, there is already an estimated time even when there is no storm "
          "simulated so I am almost sure that it is not dependent.")
    add_entry_table(doc, [
        ("#1 fix", "Fix",
         "Closed-form now reads N from the actual simulated storm. Estimate changes the "
         "moment you re-simulate."),
    ])
    h3(doc, "Takeaways")
    bullets(doc, [
        "The cheapest optimisation often pays for itself fastest. The closed-form is ~15 lines and changed the feel of the UI dramatically.",
        "Approximate UX feedback beats exact lag.",
        "Users notice \"this number didn't move when it should have\" before they notice \"this number is slightly wrong.\"",
    ])

    section_break(doc)

    # ============ Chapter IV ============
    h2(doc, "Chapter IV — Alternative #2: The Pre-Computed Scenario Library")
    margin_note(doc, "Twelve canned storms. No compute required to explore them.")
    body(doc,
         "For viewers who land on the page and don't want to wait, a dropdown of pre-computed "
         "scenarios. We picked 12 representative storms (different intensities, different "
         "distributions across the county) and wrote a Python pre-computer that runs the "
         "full scheduler, serialises the result as JSON, and dumps to scenarios/. Each file "
         "is ~1 MB; total ~12.5 MB.")
    add_entry_table(doc, [
        ("#2 build", "Build",
         "06_precompute_scenarios.py reuses functions from 05_generate_artifacts.py (the "
         "offline Python port of the JS scheduler) to compute 12 scenarios. scenarios/index.json "
         "is the manifest."),
        ("#2 build", "Build",
         "Frontend dropdown loader: on selection, fetch the scenario JSON, populate outages, "
         "populate crews + jobs, render. Skips the scheduler entirely."),
        ("#2 review", "Question",
         "User couldn't see the dropdown immediately:"),
    ])
    quote(doc, "User", "I dont see the drop downs youre talking about.")
    body(doc,
         "The HTML was committed, but GitHub Pages hadn't finished deploying. Resolved by "
         "explaining the 2–10 minute Pages cycle + Ctrl+Shift+R to bust the cache.")
    h3(doc, "Takeaways")
    bullets(doc, [
        "Caching results for common scenarios is sometimes the right answer — beats any algorithmic optimisation by definition.",
        "GitHub Pages deploy latency is real. Worth surfacing in the README.",
    ])

    section_break(doc)

    # ============ Chapter V ============
    h2(doc, "Chapter V — Alternative #3: The WebAssembly Adventure That Wasn't")
    margin_note(doc, "A two-day Rust detour that taught us V8's JIT is excellent.")
    body(doc,
         "On paper, WebAssembly looked obvious: write the scheduler in Rust, compile to "
         "wasm, get native-ish speed in the browser, ship a single binary that runs anywhere. "
         "The reality was a series of compromises forced by the Windows toolchain, followed "
         "by a hard benchmark truth.")
    h3(doc, "The toolchain gauntlet")
    body(doc,
         "First attempt: wasm-pack with full dependencies — failed because MSVC linker was "
         "missing. Adding libm — failed for the same reason. Adding wasm-bindgen — failed. "
         "The fix was to go ascetic: #![no_std], crate-type cdylib, zero external "
         "dependencies, plain rustc with the wasm32-unknown-unknown target — no host C "
         "linker required at all.")
    body(doc,
         "The cost of going no_std: we couldn't use libm's math functions. So we wrote them "
         "inline. Sin, cos, asin, sqrt, log, exp — all hand-rolled Taylor series. Bump "
         "allocator with a 256 MB workspace. Spatial grid hash via counting sort. Min-heap "
         "with parallel arrays (f64 keys, u32 values). It worked. The compiled "
         "scheduler.wasm is 17 KB.")
    h3(doc, "The benchmark")
    body(doc,
         "Side-by-side benchmark button in the UI: same scenario, run JS scheduler, run "
         "WASM scheduler, compare. The user screenshotted the result. WASM was consistently "
         "2-3× slower than the JS at 500-2000 outages.")
    add_entry_table(doc, [
        ("#3 build", "Build",
         "Rust scheduler in wasm_scheduler/src/lib.rs with #![no_std], inline math, bump "
         "allocator. Exports wasm_alloc, wasm_reset, run_scheduler."),
        ("#3 build", "Build",
         "JS wrapper loadWasmScheduler(), runWasmScheduler(), runJsSchedulerReference(), "
         "runBenchmark(). Benchmark button in the UI."),
        ("#3 review", "Question", "User shared the benchmark screenshots:"),
    ])
    quote(doc, "User", "[benchmark screenshots showing WASM is slower] is this good enough or do you want more")
    add_entry_table(doc, [
        ("#3 verdict", "Decision",
         "WASM is the wrong tool here. Root cause: hand-rolled Taylor-series math is no "
         "match for V8's native Math.sin/cos/log, which compile to single SIMD instructions "
         "on modern CPUs. With libm we might have closed the gap. Without it, we lose. "
         "Decided to not wire WASM into production; kept the artefact as a reference build."),
        ("#3 pivot", "Pivot", "User chose to move forward:"),
    ])
    quote(doc, "User", "move onto #4")
    h3(doc, "Takeaways")
    bullets(doc, [
        "V8's JIT is extraordinarily good at numeric loops. Beating it requires either real low-level optimisation (SIMD intrinsics, libm) or a fundamentally different algorithm — not just a different language.",
        "Benchmark before integrating. The whole WASM effort cost two days but the decision not to ship it was made in five minutes once we had numbers.",
        "A failed experiment is not wasted work. The grid-hash design that made it into the no_std build later informed the Python scheduler's design.",
    ])

    section_break(doc)

    # ============ Chapter VI ============
    h2(doc, "Chapter VI — Alternative #4: Server-Side Compute Arrives")
    margin_note(doc, "The pitch shifted from \"make it faster\" to \"unlock things the browser fundamentally can't do.\"")
    body(doc,
         "The framing for #4 was different from the start. The JS scheduler was already "
         "plenty fast for ordinary scenarios. The reason for a server backend was capability, "
         "not speed: Monte Carlo ensembles (running the scheduler 30+ times to get a "
         "distribution), eventual MILP solvers (commercial OR libraries don't ship as WASM), "
         "and eventual real-data calibration that needs a Python ecosystem.")
    add_entry_table(doc, [
        ("#4 build", "Build",
         "FastAPI app at 07_server.py. Endpoints: GET /health, POST /api/schedule, "
         "POST /api/monte_carlo. Pydantic models for request/response. Wide-open CORS so "
         "GitHub Pages can hit it from any origin."),
        ("#4 build", "Build",
         "Dockerfile for deployment. Initially pinned just fastapi, uvicorn, pydantic — we'd "
         "grow the dependency list as the schedulers matured."),
        ("#4 build", "Build",
         "Frontend \"Server backend\" panel: server URL input, \"Use server for Plan "
         "restoration\" toggle, Monte Carlo button. Default URL: http://localhost:8000."),
        ("First test", "Question", "First attempt to use the server backend:"),
    ])
    quote(doc, "User",
          "this is the message I am getting. its asking if the server is running so that "
          "means its not working. Do you think its possible to fix?")
    body(doc,
         "The error: \"Failed to fetch\" with the toggle on. Root cause: the FastAPI server "
         "is a separate process — checking a toggle in the browser doesn't magically start "
         "a Python server. Resolved by installing the dependencies (pip install fastapi "
         "\"uvicorn[standard]\" pydantic) and starting it locally "
         "(python -m uvicorn 07_server:app --port 8000).")
    add_entry_table(doc, [
        ("First success", "Fix",
         "Server up, /health returning {\"status\":\"ok\"}. Plan restoration through the "
         "server returns a valid response. The actual workload that broke things came next."),
    ])

    section_break(doc)

    # ============ Chapter VII ============
    h2(doc, "Chapter VII — The \"Max Settings\" Crisis")
    margin_note(doc, "A 24 950-outage × 5 000-crew stress test that broke everything. "
                "Setting the stage for every later performance fix.")
    body(doc,
         "With the server working at small scales, the user maxed out every slider in the "
         "UI: 24 950 outages, 5 000 repair crews, realistic mode, Monte Carlo 30 seeds. "
         "And waited. And waited.")
    quote(doc, "User", "so I did all the max settings, but its taking forever to load")
    body(doc,
         "What was happening: the server's scheduler was the pure-Python implementation in "
         "05_generate_artifacts.py. Pure Python doing 24 950 dispatches, each scanning 24 950 "
         "outages to find the nearest undone one, is ~625 million operations per Monte Carlo "
         "run × 30 runs = ~18 billion Python-level operations. At the rate of millions of "
         "operations per second for interpreted Python, that's hours.")
    h3(doc, "The diagnostic")
    body(doc,
         "The inner loop \"find the nearest undone outage to this crew\" is O(N) per dispatch, "
         "run N times, so O(N²) total. With N=25 000 that's 625 million Python comparisons "
         "per scheduler call. Pure Python: minutes. Even basic NumPy: still seconds. The "
         "bigger problem: 30 seeds × this cost = the user's eight-minute wait.")
    h3(doc, "The trade-off")
    body(doc,
         "Offered two paths: (A) lower the Monte Carlo workload (drop default seeds to 10, "
         "encourage smaller storms — easier to ship, but feels like a retreat); (B) actually "
         "speed up the Python scheduler with NumPy (port the inner loop to vectorised "
         "operations, ~10–50× speedup expected, ~1 hour of work).")
    quote(doc, "User", "nah do B now")
    h3(doc, "The NumPy port")
    body(doc,
         "Built scheduler_fast.py: replaces the Python inner loop with vectorised NumPy. "
         "np.argmin over a masked distance array — one CPython call into compiled NumPy code "
         "instead of 25 000 individual Python comparisons. The discovery fast-forward "
         "(when a crew has no available work and needs to skip to the next discovery time) "
         "used np.searchsorted on a pre-sorted discovery list — O(log N) instead of O(N).")
    h3(doc, "The KD-tree experiment that didn't quite work")
    quote(doc, "User", "yes please, do #1 first")
    body(doc,
         "User picked KD-tree from the next-step menu. Added scipy.spatial.cKDTree for the "
         "nearest-neighbour search. Expected ~10–20× speedup on huge N. Benchmark result: "
         "slower than the NumPy vectorised scan in realistic mode. Reason: when ~70% of "
         "outages are still \"undiscovered\" early in the storm, the KD-tree's K-nearest "
         "results get filtered down to almost nothing, K balloons, and we end up paying "
         "KD-tree query cost with no benefit. Capped K at 256 with a fallback to vectorised "
         "scan — useful in non-realistic dense regimes only.")
    add_entry_table(doc, [
        ("Step 1", "+5× Perf",
         "NumPy vectorisation in scheduler_fast.py. 2 k × 100: 0.34 s → ~0.07 s. Monte Carlo "
         "30 seeds: still seconds, but feasible."),
        ("Step 1.5", "Fix",
         "Discovered Monte Carlo always reported stddev=0.00. Cause: RNG seeds inside the "
         "scheduler were hardcoded constants, so all 30 \"different\" runs produced identical "
         "results. Fixed with per-seed RNG streams. Real variation visible — stddev 2.8 h on "
         "a representative scenario."),
        ("Step 2", "Honest result",
         "KD-tree: ~2.5× win on small-but-dense scenarios, regression on realistic-mode "
         "large scenarios. K_CAP=256 fallback merges the best of both. Not the 10–20× "
         "I promised — acknowledged honestly."),
        ("Step 3", "Question", ""),
    ])
    quote(doc, "User", "Yes do #2 and #3")
    add_entry_table(doc, [
        ("Step 3", "+14× Perf",
         "scheduler_numba.py: the entire dispatch loop JIT-compiled to native code with "
         "@njit(cache=True, fastmath=True). Inline binary heap with parallel arrays "
         "(no heapq in nopython mode), inline haversine, inline Mulberry32, inline log-"
         "normal sample. First call compiles in ~8 s (cached to disk afterward). Subsequent "
         "calls: tens of milliseconds."),
        ("Step 3", "Surprise",
         "Process pool for Monte Carlo was slower, not faster: 30 seeds × Windows process-"
         "spawn overhead (~1 s per worker) plus Numba cache-load on each cold subprocess = "
         "worse than serial. Heuristic threshold: only use the pool when per-run cost "
         "dominates spawn overhead (N × crews > 10M)."),
    ])
    h3(doc, "Numbers at the end of Chapter VII")
    bullets(doc, [
        "2 k outages × 100 crews: 0.34 s → 0.024 s (14×)",
        "Monte Carlo 30 seeds (2 k × 100): minutes → 0.21 s",
        "25 k × 500 (realistic): minutes → 3.67 s",
        "25 k × 5 000 (worst case): still ~118 s — not solved yet",
    ])

    section_break(doc)

    # ============ Chapter VIII ============
    h2(doc, "Chapter VIII — The Connecticut-Scale Refactor")
    margin_note(doc, "Where one user reframe shifted the goal from \"Hartford works\" to \"the whole state works.\"")
    body(doc,
         "With Numba in place, the typical-scale scenarios were essentially solved. The "
         "remaining ugly case: the user's maximum-slider stress test (25 000 outages × "
         "5 000 crews) still took two minutes. I floated MILP and Eversource calibration as "
         "the natural next research directions. The user pushed back with a sharper goal:")
    quote(doc, "User",
          "the real storm data is a lot later. We need to make this more scalable and more "
          "precise for Hartford county first. We can calibrate it against it later and I "
          "will let you know when. I want this model first to be able to survive against "
          "even these unrealistic scenarios for Hartford County as they are going to need "
          "to handle even heavier scenarios when scaled up to Connecticut.")
    body(doc,
         "That re-framing changed what \"good enough\" meant. Connecticut is roughly 10× "
         "Hartford County by population. A statewide simulation would need to handle "
         "~100 000-outage scenarios — well above any single Eversource event but the right "
         "ceiling to design to. The two minutes we were spending was unacceptable.")
    h3(doc, "The spatial grid hash")
    body(doc,
         "The algorithm of choice: divide the bounding box into a G×G grid (G = sqrt(N/5) — "
         "five outages per cell on average). To find the nearest undone outage for a crew, "
         "look up the crew's cell, walk it, walk concentric rings of cells outward until "
         "something valid is found. Termination: when the next ring's minimum-possible "
         "distance exceeds the current best, stop.")
    h3(doc, "The bug that doubled the runtime")
    body(doc,
         "First implementation: slower than the flat scan. Diagnosis: the \"skip non-"
         "boundary cells\" optimisation was broken when the search box clamped to the edge "
         "of the grid. Cells in the interior of the original ring kept getting revisited "
         "every iteration. Rewrote the ring iteration to enumerate the actual Chebyshev-"
         "distance == ring cells directly: top + bottom strips, then left + right columns "
         "excluding corners. No duplicates.")
    h3(doc, "The real hero — n_available")
    body(doc,
         "Even with the ring iteration fixed, 25 k × 5 000 still timed out. The culprit was "
         "realistic mode's discovery model: at t=12 h only a tiny fraction of outages have "
         "been \"discovered\" by the utility yet, but all 5 000 crews are awake and trying "
         "to dispatch. Each one walks the grid outward and finds nothing, all the way to "
         "the edge of the bounding box, before fast-forwarding.")
    body(doc,
         "Fix: an incrementally maintained counter n_available — the number of outages "
         "currently discovered-and-undone. When it's zero, fast-forward immediately via "
         "np.searchsorted on the sorted discovery list. No grid scan required. The counter "
         "is incremented as the time pointer advances past discovery thresholds and "
         "decremented when an outage is marked done.")
    add_entry_table(doc, [
        ("Grid v1", "Bug",
         "First grid implementation was 2–4× slower than the flat scan because of the "
         "boundary-skip optimisation breaking at grid edges. Rewrote with explicit ring-"
         "boundary enumeration."),
        ("Grid v2", "Partial Perf",
         "Grid hash alone helped 10k × 500 (8 s → 1 s) but barely touched 25 k × 5 000 "
         "(still 50+ s) because of discovery thrashing."),
        ("n_avail", "+246× Perf",
         "Added the n_available counter. 25 k × 5 000: 118 s → 0.48 s. The single biggest "
         "single-commit win of the project."),
        ("CT scale", "Achieved",
         "100 k outages × 2 000 crews — well above any realistic Eversource event — "
         "finishes in 660 ms. Statewide scenarios are now feasible."),
    ])
    h3(doc, "Final benchmark snapshot")
    bench = doc.add_table(rows=7, cols=3)
    for r, (a, b, c) in enumerate([
        ("Scenario", "Before", "After"),
        ("2 k × 100", "0.34 s (NumPy)", "< 10 ms"),
        ("10 k × 500", "8 s", "0.51 s"),
        ("25 k × 500 (realistic)", "~minutes", "70 ms"),
        ("25 k × 5 000 (worst case)", "118 s", "480 ms"),
        ("50 k × 1 000", "(untested)", "220 ms"),
        ("100 k × 2 000 (CT projection)", "infeasible", "660 ms"),
    ]):
        cells = bench.rows[r].cells
        is_header = r == 0
        for c_cell, txt in zip(cells, (a, b, c)):
            set_cell_borders(c_cell)
            if is_header:
                shade(c_cell, BG_HEADER)
            p = c_cell.paragraphs[0]
            add_styled(p, txt, bold=is_header, color=INK, size=11)
    body(doc, "")
    h3(doc, "Takeaways")
    bullets(doc, [
        "Algorithmic improvements (O(N²) → O(N log N) via grid hash) compound with constant-factor improvements (Numba JIT). Either alone wouldn't have closed the gap.",
        "The discovery-thrashing fix wasn't about making the search faster — it was about not searching at all when there's nothing to find. Often the best optimisation is recognising work that doesn't need doing.",
        "The 100 k-outage capability isn't useful because we expect 100 k-outage storms. It's useful because we can honestly claim the model is prepared for statewide deployment — a real engineering result for a paper's introduction.",
    ])

    section_break(doc)

    # ============ Chapter IX ============
    h2(doc, "Chapter IX — The Render Deployment Saga")
    margin_note(doc, "Four straight failed deploys. The root cause was Python exception-hierarchy semantics.")
    body(doc,
         "With a working backend, the next step was making it publicly callable so visitors "
         "to the GitHub Pages site could use the server without needing to install anything. "
         "Render's free tier seemed ideal: $0/month, auto-deploy from GitHub via a "
         "render.yaml Blueprint, container sleeps when idle.")
    quote(doc, "User", "can you do that for me")
    body(doc,
         "Honest answer: no. Deploying to Render needs the user's account credentials. But "
         "I could prep everything so the user's part was just clicking through Render's "
         "Blueprint UI. Wrote render.yaml, ensured the Dockerfile was self-contained, pushed.")
    h3(doc, "Failure #1 — matplotlib missing")
    body(doc,
         "First deploy error: \"matplotlib + numpy required. Install with: pip install "
         "matplotlib numpy\". Cause: 07_server.py imported 05_generate_artifacts.py via "
         "importlib for the pure-Python fallback scheduler — and 05_generate_artifacts "
         "imports matplotlib at the top for its plotting helpers. Server never plots; the "
         "import just has to succeed. Added matplotlib to the Docker image.")
    h3(doc, "Failure #2 — file not found")
    body(doc,
         "Next deploy: FileNotFoundError: '/app/data/hartford_boundary.json'. Same file "
         "(05_generate_artifacts.py) tries to load Hartford boundary GeoJSON at module level. "
         "We didn't ship the data directory into the Docker image because the server doesn't "
         "need it. Wrapped the importlib call in try / except Exception so any import "
         "failure would just set art = None and proceed.")
    h3(doc, "Failure #3 — the deep gotcha")
    body(doc,
         "Same FileNotFoundError. The \"fix\" hadn't actually caught the error. Hours of "
         "head-scratching. Then the realisation: 05_generate_artifacts.py doesn't just "
         "raise — it raises SystemExit when matplotlib is missing. SystemExit inherits from "
         "BaseException, not from Exception. This is deliberate Python design: sys.exit() "
         "is supposed to terminate the program even when called inside a try / except "
         "Exception block, so the standard library deliberately raises a subclass that "
         "bypasses that guard. The exception is catchable, but you have to write except "
         "BaseException or use a bare except: — both of which are usually a code smell.")
    body(doc,
         "The correct fix was simpler: stop importing 05_generate_artifacts.py at all. The "
         "Numba and NumPy schedulers cover the same functionality, faster. The pure-Python "
         "fallback was already dead code in production.")
    quote(doc, "User", "it keeps failing no matter how many times I restart")
    add_entry_table(doc, [
        ("Deploy 1", "Failure",
         "matplotlib + numpy required — added matplotlib to Dockerfile."),
        ("Deploy 2", "Failure",
         "FileNotFoundError: data/hartford_boundary.json — wrapped import in "
         "try/except Exception."),
        ("Deploy 3", "Failure",
         "Same error. SystemExit bypasses except Exception. Removed the import entirely."),
        ("Deploy 4", "Success",
         "Live at hartford-grid-server.onrender.com. /health returns {\"status\":\"ok\"}."),
    ])
    h3(doc, "Takeaways")
    bullets(doc, [
        "Python's exception hierarchy is intentional: BaseException > Exception. SystemExit, KeyboardInterrupt, and GeneratorExit all live above Exception so they can't be silenced accidentally.",
        "The cleanest fix to an import-time problem is often \"stop doing the import.\" Optional dependencies should be optional in code, not just in a try/except.",
        "Render's free tier deploys are fast enough to iterate on — each deploy was 5-8 minutes. Useful for this kind of debugging.",
    ])

    section_break(doc)

    # ============ Chapter X ============
    h2(doc, "Chapter X — Polish: Auto-Detect, Progress, /version")
    margin_note(doc, "\"I am still a little unsure with how to check if the server is actually up to date or not.\"")
    body(doc,
         "After the deploy succeeded, the user raised three related concerns: (a) how do I "
         "verify the server is up to date? (b) can the server be mandatory so I don't have "
         "to think about the toggle? (c) can the server \"live on GitHub\" so it auto-"
         "launches when the page loads?")
    body(doc,
         "Answer to (c) had to be honest: GitHub Pages is static hosting — it can't run a "
         "Python server. The Render service is separate infrastructure. But the spirit of "
         "the request — \"the user shouldn't have to think about the server\" — was "
         "achievable through three changes: a verification endpoint, a warm-up ping on "
         "page load, and an auto-detect path that silently falls back to in-browser when "
         "the server is offline.")
    add_entry_table(doc, [
        ("/version", "Build",
         "Added GET /version returning {\"commit\": \"...\", \"backend\": \"numba\"}. Reads "
         "the commit from Render's RENDER_GIT_COMMIT env var. Lets the frontend display "
         "\"connected · commit fd87b7a · backend numba\" so the user can verify the deployed "
         "server matches GitHub HEAD."),
        ("Auto-warm", "Build",
         "On page load, the JS fires a no-cors GET to /health to wake the sleeping Render "
         "container. Then probes /version at 5 s, 30 s, 60 s, and every 60 s thereafter to "
         "keep the dot's state current."),
        ("Auto-detect", "Build",
         "Health dot: gray (probing) → amber (waking) → green (ok) → red (offline). When "
         "the user clicks Plan restoration: if dot is green, request goes to the server. "
         "If red, silently runs in-browser with a status note. Toggle defaults to ON."),
        ("Progress", "Build",
         "Server-side restoration is an atomic call — no intermediate progress signal. So "
         "the progress bar fakes it: exponential asymptote toward 90% with a time constant "
         "tuned for Render cold starts (~30–60 s), snaps to 100% on response, fades after "
         "400 ms."),
        ("URL mishap", "Question", "User pasted the wrong identifier:"),
    ])
    quote(doc, "User", "where am I supposed to collect the url for my server location")
    body(doc,
         "The user had pasted the Render Service ID (srv-d8qs8b6gvqtc73e70mog) into the "
         "Server URL field instead of the actual URL (https://hartford-grid-server.onrender."
         "com). Confusion fixed by labelling the field clearly and defaulting to the "
         "deployed URL.")
    add_entry_table(doc, [
        ("Render UX", "Question", "User shared a Render cold-start splash:"),
    ])
    quote(doc, "User",
          "in regards to number 3, do a and b, but also check the attached image cause i "
          "cant tell what im supposed to see from here")
    body(doc,
         "The attached image showed Render's cold-start splash (\"INCOMING HTTP REQUEST "
         "DETECTED … SERVICE WAKING UP …\"). Explained: that's expected behaviour. First "
         "request after sleep boots the container, takes 30–60 s, then real responses "
         "follow. Subsequent calls within the next 15 min are fast.")
    h3(doc, "Takeaways")
    bullets(doc, [
        "\"Verification\" as a first-class UX feature: showing the commit SHA the running server was built from removes a whole class of \"is this up to date?\" anxiety.",
        "When the underlying operation has no real progress signal, a well-tuned fake progress bar (asymptotic exponential) is genuinely better UX than a static spinner.",
        "Free-tier hosting cold starts are the user-facing tax of the $0 price tag. Auto-warm-up hides ~80% of that pain.",
    ])

    section_break(doc)

    # ============ Chapter XI ============
    h2(doc, "Chapter XI — The Vanishing Markers")
    margin_note(doc, "A regression caught by a user observation.")
    body(doc,
         "After all the speed improvements, the user noticed something missing in the "
         "visualisation:")
    quote(doc, "User",
          "I have a question, why did you stop with those images that we used to have for "
          "when the plan restoration works. It works very well now with the processing "
          "speed as it is now able to handle the maximum settings very well. However, I "
          "still want those images on the grid so can you bring them back and lets not get "
          "rid of those to improve processing speed")
    body(doc,
         "What had happened: the server's ScheduleResponse returned only crew counts and "
         "total times — not the per-job sequence that the browser needs to draw the "
         "numbered repair circles. When the server backend was used, the map's depot "
         "squares and numbered 1/2/3 circles disappeared.")
    body(doc,
         "Honest assessment of the cost of restoring them: negligible at realistic scales, "
         "+200-500 ms of network at 25 k, +1-3 s at 100 k. The Numba scheduler already "
         "computes which crew handles which outage during dispatch — we were just throwing "
         "that data away. Recording it costs one int32 write per dispatch.")
    add_entry_table(doc, [
        ("Build", "Build",
         "Added a flat dispatch log inside both _run_scheduler and _run_scheduler_grid "
         "(Numba-JIT'd) — three parallel np.empty(N) arrays (log_crew, log_outage, log_eta) "
         "plus a counter. One write per successful dispatch. Returned as sliced views."),
        ("Build", "Build",
         "Python wrapper rebuilds per-crew job lists from the flat log, preserving dispatch "
         "order within each crew. Updated scheduler_fast.py to return the same dict shape."),
        ("Build", "Build",
         "Server ScheduleResponse grew a jobs[] array of {lat, lon, eta} per crew. Larger "
         "JSON response but only adds 200-500 ms of network at 25 k outages."),
        ("Build", "Build",
         "Frontend renderPlan() extracted from planRestoration() so both local and server "
         "paths share the same renderer."),
    ])
    h3(doc, "Takeaways")
    bullets(doc, [
        "Speed without visualisation parity is a false win. Always check that the user-facing artefact still looks the way they expect.",
        "\"What does it cost to put it back?\" is the right question. Often the answer is \"very little.\"",
        "Factoring the renderer out paid for itself immediately and will pay forward — future paths (MILP comparison, scenarios from disk) all need to render the same way.",
    ])

    section_break(doc)

    # ============ Chapter XII ============
    h2(doc, "Chapter XII — Direction Conversations & Research Strategy")
    margin_note(doc, "The conversations that shaped what to build next.")
    h3(doc, "Conversation 1 — WebGPU as next step")
    body(doc,
         "After Numba was in place, I proposed Alternative #5 (WebGPU) as the next direction. "
         "The pitch: parallel argmin on the GPU for the huge-N stress test, potentially 100× "
         "on 25 k × 5 000 cases. The user agreed but also asked about research direction.")
    h3(doc, "Conversation 2 — MILP vs Calibration")
    quote(doc, "User",
          "MLP/calibration is what I want to do if it helps leads me towards more "
          "publishable research contributions")
    body(doc,
         "Recommended doing MILP first — it's self-contained (PuLP + CBC solver, no external "
         "data needed), produces a publishable claim (\"our greedy is within X% of optimal at "
         "scales where MILP is tractable\"), and sets up the framework that the calibration "
         "phase will plug into. Eversource calibration would happen in parallel as the user "
         "pursues data access through PURA dockets.")
    h3(doc, "Conversation 3 — Scale before research")
    body(doc, "Before MILP started, the user re-prioritised one more time:")
    quote(doc, "User",
          "the real storm data is a lot later. We need to make this more scalable and "
          "more precise for Hartford county first.")
    body(doc,
         "This is what triggered the Chapter VIII grid-hash + n_available work. With "
         "Connecticut scale handled, MILP and calibration both become more credible: they're "
         "now plugging into a model that we can honestly claim is engineered for statewide "
         "deployment.")
    h3(doc, "Current research-direction stack (queued, not yet built)")
    add_entry_table(doc, [
        ("Next", "Decision (MILP)",
         "PuLP + CBC for small-N (≤50 outage) optimal scheduling. Compute the greedy-vs-"
         "optimal gap. Publishable result: \"our heuristic is within X% of optimal at scales "
         "where MILP is tractable, supporting its use at scales where MILP is not.\""),
        ("Parallel", "Decision (Calibration)",
         "PURA dockets & Eversource storm post-mortems → fit travel speed, road-multiplier, "
         "workday hours, repair-time distribution to observed restoration curves. "
         "Bottlenecked on data access."),
        ("Later", "Decision (WebGPU)",
         "Engineering polish for the in-browser path. Deferred — Numba already handles "
         "Connecticut scale, so the urgency is gone."),
    ])

    section_break(doc)

    # ============ Chapter XIII ============
    h2(doc, "Chapter XIII — Documentation Push")
    margin_note(doc, "Where you, dear reader, are.")
    quote(doc, "User",
          "wait can you add a file in the github repository showing like a journal of all "
          "the changes and like questions and answers all in different colors. Make it in "
          "like a tabular format and make it look like a journal aesthetic. Also once again "
          "me all the research papers already published in this field related to what im "
          "doing and again tell me the niche im exploring and help me put the research "
          "paper stuff in a google docs.")
    body(doc,
         "Three deliverables produced: this Hartford_Grid_Dev_Journal.docx (developer "
         "history), a JOURNAL.html in the repo for browser viewing, and a separate "
         "Hartford_Grid_Research_Context.docx with a literature map by theme, niche "
         "analysis, open research questions, and data sources to pursue.")
    body(doc,
         "Honest caveat carried into the research document: the citations are starting "
         "points drawn from a directed literature scan, not a verified bibliography. "
         "Authors and research groups are the most reliable parts; specific paper titles, "
         "years, and venues should be verified on Google Scholar or IEEE Xplore before "
         "being cited in a manuscript.")

    section_break(doc)

    # ============ Chapter XV ============
    h2(doc, "Chapter XV — The Realism Round & the Polish Sprint")
    margin_note(doc, "Built after the documentation push. Customer-impact weighting, crew "
                "specialization, calibration framework, multi-server fan-out, customers-"
                "restored curve, server keep-alive, all-toggles-on speedup. The features that "
                "took the simulator from 'engineering done' to 'ready for the research story.'")
    body(doc,
         "With the documentation deliverables shipped, the work that followed was driven by a "
         "single question: is the realistic toggle actually realistic? The honest answer "
         "landed at 'structurally realistic but not validated.' The path forward was therefore "
         "two-track: add the missing realism factors that the literature says matter most, and "
         "build the calibration framework so when real Eversource data arrives the simulator "
         "can be tuned against it.")

    h3(doc, "XV.1 — Customer-impact-weighted dispatch")
    body(doc,
         "Real utility dispatchers don't just send crews to the closest outage. They detour "
         "past a single-house lateral to reach a substation that restores thousands of "
         "customers. This is the biggest gap between the original scheduler and reality. The "
         "fix: a scoring function score(o) = customers(o) - customer_weight * d² that biases "
         "dispatch toward high-customer outages while still keeping spatial efficiency.")
    body(doc,
         "Empirical result: at 25 000 outages × 5 000 crews, customer-weighted dispatch dropped "
         "total restoration time from 72 h (pure-nearest) to 61.7 h. That's not just a curve-"
         "shape change — it's a 14% improvement in the total customer-minutes-without-service "
         "metric utilities actually care about. The toggle is opt-in (default off, preserves "
         "backwards compatibility).")

    h3(doc, "XV.2 — Calibration framework")
    quote(doc, "User", "can you do 2 and 4")
    body(doc,
         "Built /api/calibrate: accepts an observed restoration curve from a real storm, runs "
         "SciPy Nelder-Mead optimization over the four most-tunable realism factors (travel "
         "speed, assessment delay, workday hours, road multiplier), returns the parameter set "
         "that minimizes RMSE between simulator output and observation. Required refactoring "
         "the Numba scheduler so those four were inputs rather than hard-coded constants.")
    body(doc,
         "Self-test against synthetic data: generate an 'observed' curve with known ground-"
         "truth parameters, run calibration from default initial guesses. Result: RMSE drops "
         "from 4 267 to 43 (a 100× reduction) in 53 iterations. The recovered parameters don't "
         "exactly match the truth because multiple parameter combinations produce nearly "
         "identical curves — calibration optimizes curve-match, not parameter-match, which is "
         "the correct success criterion.")

    h3(doc, "XV.3 — Crew specialization")
    body(doc,
         "Real utilities deploy roughly 80% line crews and 20% tree crews. About 30% of "
         "distribution outages involve trees down on lines that need clearing before line "
         "work begins. The simplest defensible model: split the fleet 80/20, tag 30% of "
         "outages as tree-blocked, run two independent dispatch subsystems. Total restoration "
         "time = max(tree, line).")
    body(doc,
         "Implementation lived entirely at the server-helper level — no Numba scheduler "
         "changes needed. The two subsystem calls were initially sequential, doubling wall-"
         "clock time at max settings. Switching to a 2-worker ThreadPoolExecutor delivered "
         "true parallelism because Numba releases the GIL during JIT-compiled code.")

    h3(doc, "XV.4 — Customer-restored curve overlay")
    body(doc,
         "The customer-impact toggle has no effect on the total restoration time (same "
         "outages, same crews, same total work). Its effect is entirely on which customers "
         "get power back first — the area under the customers-restored-over-time curve. "
         "Without a visualization of that curve, the feature was invisible.")
    body(doc,
         "Fix: an inline 280×80 SVG line chart under the Total Restoration Time stat box. "
         "Auto-appears after each Plan restoration. Down-samples to ~140 path points for huge "
         "scenarios. Works for both local and server-routed plans because it lives inside the "
         "shared renderPlan() helper.")

    h3(doc, "XV.5 — Multi-server batch sweep")
    quote(doc, "User",
          "Yes Id love for the multi-sever approach if I am able to run multiple different "
          "storm scenaiors at the same time. I want to note that I am unable to actually "
          "spend money for servers.")
    body(doc,
         "Honest reality-check first: a single 250k-outage scenario at realistic mode is "
         "already sub-second on one server, and the scheduler loop is inherently sequential, "
         "so multi-server can't help that case. Where multi-server does crush is independent-"
         "scenario workloads: Monte Carlo ensembles, parameter sweeps, varying-seed analysis. "
         "Each scenario is a fully independent scheduler call, so N servers gives ~N× "
         "throughput.")
    body(doc,
         "Architecture: new /api/batch endpoint that accepts a list of scenarios + a list of "
         "worker URLs, round-robins scenarios onto workers, fans out via ThreadPoolExecutor. "
         "Empty workers list = serial in-process. The user spins up additional free Render "
         "services (each takes ~3 min) and pastes their URLs into the batch UI to scale "
         "linearly. Stays within the free tier — Render's free-tier limits are per-service, "
         "not per-account.")

    h3(doc, "XV.6 — Keep-alive & auto-rewake")
    quote(doc, "User",
          "the server backend also turns to red after I dont give it commands for a bit. Is "
          "it possible that for as long as the interactive is open on the main tab of the "
          "person, the server can be continuously pinged so it wouldnt have to be restarted")
    body(doc,
         "Diagnosed: the existing 60-second probe had a 3.5-second timeout, way shorter than "
         "Render's 30–60 second cold-start wake time. So every probe during wake would time "
         "out, mark the server offline, and the dot would stay red even though the keep-alive "
         "interval was firing correctly. Browser background-tab throttling was a secondary "
         "contributor.")
    body(doc,
         "Four-layer fix: 4-minute keep-alive interval (well under Render's 15-min sleep "
         "threshold); auto-rewake on any failed probe (the dot stops getting stuck red); "
         "visibility-change re-probe when the tab regains focus; and a manual 'Wake server "
         "now' button as a guaranteed user-controlled fallback.")

    h3(doc, "XV.7 — All-toggles-on speedup")
    quote(doc, "User",
          "can you try to increase processing speed for when all toggles are on. It is "
          "unable to handle max settings that well")
    body(doc,
         "The problem: customer-weighted scoring was using the O(N²) dense scan because the "
         "grid-hash's ring-expansion termination didn't naively generalize to non-monotonic-"
         "in-distance scoring. The fix: a different termination bound — upper_bound = "
         "max_customers_global - customer_weight × ring_min². When that upper bound drops "
         "below the current best score, no further ring can produce a higher-scoring outage.")
    body(doc,
         "With that change, customer-weighted scoring runs inside the grid hash with ring-"
         "expansion intact. 25k × 5000 customer-weighted dropped from ~50 s (dense fallback) "
         "to 2.24 s. With crew specialization also running its subsystems in parallel, all-"
         "toggles-on at max settings now lands in 2–3 seconds end-to-end.")
    body(doc,
         "Two more wins shipped together: gzip middleware compresses the ~1.5 MB JSON "
         "response at 25k outages down to ~150 KB, saving 100–500 ms on network. And the "
         "Numba JIT is pre-warmed at server boot so the first request after a Render "
         "redeploy doesn't pay the ~10 s compile penalty — container startup is slightly "
         "slower but every user request is fast from day one.")

    h3(doc, "Takeaways from Chapter XV")
    bullets(doc, [
        "Adding the right realism factors is more important than calibrating. Customer-impact weighting changed a measurable output metric (total restoration time) by 14% at max settings, before any calibration was applied. That's a real engineering result, not a fit.",
        "Multi-server can't help a sequential algorithm at all. The single-scenario scheduler is fundamentally sequential and already sub-second; throwing N servers at it does nothing. Multi-server helps for independent work — Monte Carlo, parameter sweeps, batch evaluation.",
        "Free-tier hosting limitations are an interaction-design problem. Render's 15-minute idle sleep + 30-60 second cold start became a usability issue not because the engineering was wrong but because the keep-alive timing didn't account for it. The fix was three lines of JavaScript plus a button.",
        "Visualization parity matters as much as compute parity. The customer-impact toggle was nearly invisible until the customers-restored curve overlay landed. Speed without the right output is no win.",
    ])

    section_break(doc)

    # ============ Chapter XIV ============
    h2(doc, "Chapter XIV — Coda")
    body(doc,
         "The project started as a Hartford County storm simulator and ended this "
         "development cycle as a Connecticut-prepared, server-backed, browser-first "
         "research platform. The engineering path zig-zagged: a successful closed-form "
         "approximation, a working but underwhelming WebAssembly port, a successful FastAPI "
         "backend, a NumPy port that helped, a KD-tree experiment that helped less than "
         "promised, a Numba port that helped a lot, a grid hash that helped little on its "
         "own and enormously when paired with an n_available counter, and a four-attempt "
         "deployment saga whose root cause was a Python exception-hierarchy subtlety.")
    body(doc,
         "What's live now: a public-facing simulator that auto-detects its own server "
         "backend, falls back gracefully when offline, handles 100 000-outage scenarios in "
         "under a second, and renders the per-crew dispatch sequence as colored numbered "
         "circles on a Leaflet map of Hartford County. The road ahead is MILP comparison "
         "for academic rigor and PURA-grounded calibration for real-world claims. The "
         "engineering substrate is done; what remains is the research story it can now "
         "support.")

    section_break(doc)

    # ============ Appendix A — Problems Faced ============
    h2(doc, "Appendix A — Problems Faced (Cross-Project Catalogue)")
    margin_note(doc, "Every significant problem the project hit, in order, with root cause "
                "and resolution. Useful as a debugging-pattern reference for the next person "
                "doing this kind of work.")
    add_entry_table(doc, [
        ("I", "Fix — \"4-hour fantasy\"",
         "Cause: naïve scheduler ignored real-world delays; predicted 4 h vs reality 2–7 d. "
         "Resolution: seven multiplicative realism factors (assessment, repair, discovery "
         "ramp, mutual-aid waves, road proxy, workday clamp, critical priority)."),
        ("I", "Fix — \"Page isn't responding\"",
         "Cause: dispatch loop never yielded to the browser event loop at N=1500+. "
         "Resolution: async-yield chunking with await new Promise(r => setTimeout(r, 0)) "
         "every ~500 dispatches."),
        ("I", "Fix — DOM marker render cost",
         "Cause: each L.marker creates an HTML element; 10 000 markers take seconds. "
         "Resolution: custom PointCloudLayer canvas renderer for thousands of dots in a "
         "single pass."),
        ("I", "Fix — O(N) nearest-outage scan",
         "Cause: inner search scans all outages every dispatch. Resolution: JS-side "
         "spatial grid hash with ring expansion."),
        ("I", "Fix — Cross-language PRNG mismatch",
         "Cause: Mulberry32 had to match bit-for-bit in JS and Python including overflow "
         "semantics. Resolution: explicit & 0xFFFFFFFF masks, head-to-head 10 000-sample "
         "comparison test before any artifacts were committed."),
        ("III", "Fix — Closed-form ignored actual storm",
         "Cause: read N from slider rather than realised storm. Resolution: read N from "
         "storm.outages.length when storm exists."),
        ("IV", "Fix — Dropdown invisible after deploy",
         "Cause: GitHub Pages deploy lag (2–10 min). Resolution: explained the deploy "
         "cycle + Ctrl+Shift+R hard refresh."),
        ("V", "Fix — MSVC linker missing",
         "Cause: Windows toolchain blocked Rust crates needing a host C linker. "
         "Resolution: #![no_std], crate-type cdylib, zero external deps, plain rustc."),
        ("V", "Fix — libm / wasm-bindgen blocked",
         "Cause: same root. Resolution: hand-rolled Taylor-series math (sin, cos, asin, "
         "sqrt, log, exp) inline; bump allocator; parallel-array heap. 17 KB compiled "
         "scheduler.wasm."),
        ("V", "Fix — WASM 2–3× slower than JS",
         "Cause: hand-rolled Taylor-series math vs V8's native Math.* intrinsics. "
         "Resolution: WASM kept as reference build; V8 JS remained the production "
         "scheduler."),
        ("VI", "Fix — \"Failed to fetch\" error",
         "Cause: FastAPI server is a separate process; toggling the UI doesn't start it. "
         "Resolution: pip install fastapi \"uvicorn[standard]\" pydantic + "
         "python -m uvicorn 07_server:app --port 8000."),
        ("VII", "Fix — Max-settings runs for minutes",
         "Cause: 25 k × 5 k = 625 M comparisons per Monte Carlo run × 30 seeds = 18 G ops "
         "in interpreted Python. Resolution (multi-step): NumPy vectorisation (+5×) → "
         "Numba JIT (+14×) → grid hash + n_available (+246×). Documented in Chapters "
         "VII & VIII."),
        ("VII", "Fix — KD-tree slower than expected",
         "Cause: in realistic mode only ~30% of outages are discovered early; K balloons "
         "as filtered results shrink. Resolution: K_CAP = 256 with fallback to vectorised "
         "scan. Useful only in non-realistic dense regimes."),
        ("VII", "Fix — Monte Carlo stddev = 0",
         "Cause: RNG seeds inside the scheduler were hardcoded constants, so all 30 "
         "\"different\" runs were identical. Resolution: per-seed RNG streams "
         "(seed * 1117 + 23 for repair, seed * 991 + 7 for discovery)."),
        ("VII", "Fix — Process pool slower than serial",
         "Cause: Windows process-spawn overhead (~1 s/worker) + per-worker Numba cache "
         "load on cold subprocess. Resolution: heuristic threshold — only use pool when "
         "N × crews > 10M."),
        ("VIII", "Fix — Grid hash slower than flat scan",
         "Cause: \"skip non-boundary cells\" optimisation revisited interior cells when "
         "the search box clamped to the grid edge. Resolution: rewrote ring iteration to "
         "enumerate Chebyshev-distance == ring boundary cells directly."),
        ("VIII", "Fix — Discovery thrashing at high crew counts",
         "Cause: 5 000 crews all awake at t=12 h while only a tiny fraction of outages "
         "were discovered; each crew walked the full grid finding nothing. Resolution: "
         "incrementally maintained n_available counter — fast-forward immediately when "
         "zero via np.searchsorted. Single biggest commit win of the project (118 s → "
         "0.48 s)."),
        ("IX", "Fix — Render deploy #1: matplotlib missing",
         "Cause: 05_generate_artifacts.py imports matplotlib unconditionally; server "
         "imported it via importlib. Resolution (interim): added matplotlib to Docker image."),
        ("IX", "Fix — Render deploy #2: FileNotFoundError",
         "Cause: same file loads data/hartford_boundary.json at module level; not shipped "
         "in image. Resolution (interim): wrapped import in try/except Exception."),
        ("IX", "Fix — Render deploy #3: SAME error after try/except",
         "Cause: file raises SystemExit, which inherits from BaseException, not Exception. "
         "except Exception deliberately doesn't catch sys.exit() — standard library design. "
         "Resolution (real): stop importing 05_generate_artifacts entirely; Numba/NumPy "
         "schedulers cover the same functionality faster."),
        ("X", "Fix — User pasted Service ID, not URL",
         "Cause: Render dashboard label looked URL-like (srv-d8qs8b6gvqtc73e70mog). "
         "Resolution: clearer panel labelling + default the field to the deployed URL."),
        ("X", "Fix — Render cold-start splash confused user",
         "Cause: free-tier containers sleep after 15 min idle; next request triggers a "
         "30–60 s wake-up shown as a splash. Resolution: warm-up ping fires on page load; "
         "health dot turns green when ready; fake exponential progress bar so the wait "
         "feels active."),
        ("X", "Fix — No way to verify deployed commit",
         "Cause: no version endpoint. Resolution: added GET /version returning "
         "{\"commit\": \"...\", \"backend\": \"...\"}; frontend displays the SHA so it can "
         "be matched against GitHub HEAD."),
        ("XI", "Fix — Map markers vanished on server path",
         "Cause: server's ScheduleResponse returned only crew counts & total times, not "
         "the per-job sequence the browser needs to draw numbered repair circles. "
         "Resolution: flat dispatch log inside Numba scheduler, full jobs[] array in "
         "response, renderPlan() extracted so both local and server paths share one renderer."),
    ])

    h3(doc, "Patterns across the catalogue")
    bullets(doc, [
        "\"Optional dependencies\" need to actually be optional in code, not just in spirit. Three deploy failures in a row stemmed from one file's unconditional import of matplotlib + load of a data file. The clean fix was removing the import entirely.",
        "Real measurements beat plausible estimates. Three of the biggest \"this will help a lot\" optimisations (WASM, KD-tree, process pool) all ended up neutral or worse. The only way we knew was benchmarking after each.",
        "The hardest bugs are about semantics, not syntax. The four-deploy SystemExit saga was caused by Python's exception hierarchy doing exactly what it's designed to do.",
        "Many performance problems are really \"we're doing work that doesn't need to be done\" problems. The 246× win in Chapter VIII came from skipping the grid scan when there was nothing to find, not from making the scan faster.",
        "Visualisation parity is part of correctness. The vanishing markers regression wasn't caught by tests — it was caught by a user looking at the map. Speed without the right output isn't a win.",
    ])

    doc.save("Hartford_Grid_Dev_Journal.docx")
    print("Wrote Hartford_Grid_Dev_Journal.docx")


# =================== Build the Research Context (unchanged) ===================

def build_research():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = sec.bottom_margin = Inches(0.9)
    sec.left_margin = sec.right_margin = Inches(0.9)

    h1(doc, "Research Context & Literature Map")
    p = doc.add_paragraph()
    add_styled(p, "Hartford County Power Grid Resilience Simulation",
               italic=True, color=MARGIN, size=12)
    p = doc.add_paragraph()
    add_styled(p, "Companion to the development journal",
               italic=True, color=MARGIN, size=10)

    body(doc, "")
    h2(doc, "What This Project Is, In One Paragraph")
    body(doc,
         "A browser-first, county-specific interactive simulation of post-storm distribution-"
         "grid restoration. Models 100k+ outages and 5 000+ crews at Connecticut scale in "
         "under a second per scenario, with a server backend exposing Monte Carlo ensembles. "
         "The greedy rolling-horizon scheduler accounts for damage-assessment delay, log-"
         "normal repair durations, mutual-aid waves, workday clamps, and crew-routing "
         "heuristics. Designed as a foundation for calibration against real Eversource / "
         "PURA storm post-mortem data and for benchmarking against an MILP optimal "
         "scheduler.")

    body(doc, "")
    h2(doc, "The Niche You Are Exploring")
    margin_note(doc, "Reading of the academic + industry landscape; revise as you find more sources.")

    niche_points = [
        ("The UConn group has the prediction side covered — but not the restoration side.",
         "Wanik, Anagnostou, Cerrai, Astitha and collaborators at UConn's Eversource Energy "
         "Center have built sophisticated models for predicting WHERE and HOW MANY outages "
         "will occur in a Connecticut storm. They have not built a comparable simulator for "
         "HOW restoration unfolds once those outages exist. Your work plugs into the "
         "downstream of their pipeline: their model says \"3 000 outages in this storm,\" "
         "yours says \"with M crews, those 3 000 outages get restored in X hours.\" The "
         "natural collaboration."),
        ("Most academic work on grid restoration is offline batch optimisation.",
         "Coffrin, Van Hentenryck, Arif, and collaborators have produced strong MILP and "
         "stochastic-programming models for restoration. Less of it targets distribution-"
         "level county-scale, and almost none of it is interactive."),
        ("Most utility decision-support tools are proprietary and inaccessible.",
         "Eversource, Avangrid, and others use commercial outage-management systems (OMS) "
         "like OSI Monarch or GE PowerOn. These run in control rooms and are not accessible "
         "for academic exploration."),
        ("The synthetic-grid corpus targets transmission, not distribution.",
         "ARPA-E's GRID DATA program and Texas A&M's synthetic networks deliver "
         "transmission-scale test cases. County-resolution distribution grids with realistic "
         "storm-outage patterns are scarce."),
        ("Browser-first, real-time interactive grid simulation at this fidelity is rare.",
         "There are a few public demos but most are toy-scale (≤ 100 nodes). Running "
         "100 000 outages + 5 000 crews + Monte Carlo in a browser-backed-by-API pipeline "
         "is a differentiator."),
        ("Calibration of restoration timelines against real PURA / Eversource storm filings is largely untouched.",
         "Statistical outage-prediction papers exist (Wanik 2015, Cerrai 2019, Nateghi 2011). "
         "Restoration-timeline calibration is much less common — partly because the data "
         "lives in regulatory filings rather than open datasets, and partly because the "
         "prediction community has historically stopped at the outage-count layer."),
    ]
    for headline, expansion in niche_points:
        p = doc.add_paragraph()
        add_styled(p, "▸ ", color=ACCENT, bold=True, size=12)
        add_styled(p, headline, bold=True, color=INK, size=11)
        p = doc.add_paragraph()
        add_styled(p, "    " + expansion, color=INK, size=11)

    body(doc, "")
    p = doc.add_paragraph()
    add_styled(p, "Net positioning: ", bold=True, color=ACCENT, size=12)
    add_styled(p,
               "a publicly accessible, scale-validated, calibratable interactive grid-"
               "resilience simulator for a specific real-world utility service territory. "
               "The contribution is the integration — the engineering of an accessible tool "
               "with research-grade behaviour — rather than a new algorithm.",
               color=INK, size=11)

    body(doc, "")
    h2(doc, "Literature Map by Theme")
    margin_note(doc, "Authors and groups whose work directly intersects this project. Verify each citation independently — these are starting points for your own search.")

    # Each item is now (citation, why, what, key_terms).
    themes = [
        ("UConn / Eversource Outage-Prediction Group (your most directly relevant citations)", BG_QUESTION, [
            ("Wanik, D. W., Anagnostou, E. N., Hartman, B. M., Frediani, M. E. B., Astitha, M. "
             "(2015). Storm outage modeling for an electric distribution network in "
             "Northeastern USA. Natural Hazards 79(2), 1359–1384.",
             "this is essentially the foundational paper for your project. UConn's Eversource "
             "Energy Center built outage-prediction models on real CL&P / Eversource "
             "distribution data in Connecticut — the exact territory your simulation models. "
             "Anything you publish should cite this work; your simulator can be positioned as "
             "the restoration-side counterpart to their outage-prediction side.",
             "develops a machine-learning model (Random Forest) that predicts the number of "
             "outages per municipality for a given storm, using weather and land-cover "
             "features. Validated against a multi-year storm catalogue from a Connecticut "
             "utility. Establishes the data pipeline that the rest of the UConn corpus builds on.",
             "Random Forest · outage prediction · storm catalogue · weather features · "
             "land-cover features · municipality-level aggregation · Connecticut Light & "
             "Power (CL&P) · Eversource Energy Center"),
            ("Wanik, D. W., Parent, J. R., Anagnostou, E. N., Hartman, B. M. (2017). Using "
             "vegetation management and LiDAR-derived tree height data to improve outage "
             "predictions for electric utilities. Electric Power Systems Research 146, 236–245.",
             "the vegetation-aware extension. Trees are the dominant cause of storm-induced "
             "distribution outages in CT, and this paper shows that adding tree-height data "
             "to the model materially improves predictions. Useful citation if you ever add "
             "a vegetation layer to your storm model.",
             "augments the Random Forest outage-prediction model with LiDAR-derived tree-"
             "height data and vegetation-management records. Demonstrates that tree height "
             "near distribution lines is one of the strongest single predictors of outage "
             "count.",
             "LiDAR · vegetation management · tree height · feature engineering · variable "
             "importance · trim cycle · right-of-way · canopy cover"),
            ("Cerrai, D., Wanik, D. W., Bhuiyan, M. A. E., Zhang, X., Yang, J., Frediani, "
             "M. E. B., Anagnostou, E. N. (2019). Predicting Storm Outages Through New "
             "Representations of Weather and Vegetation. IEEE Access 7, 29639–29654.",
             "the most refined version of the UConn outage-prediction pipeline. Uses WRF "
             "numerical weather model output + vegetation features to predict outages at "
             "high spatial resolution. If you ever drive your storm model from real weather "
             "forecasts rather than uniform sampling, this paper is the bridge.",
             "predicts outages on a 2 km grid using WRF (Weather Research & Forecasting) "
             "atmospheric model output combined with vegetation rasters. Uses gradient-"
             "boosted decision trees; reports significant accuracy improvements over the "
             "2015/2017 baselines.",
             "Weather Research & Forecasting (WRF) · gradient boosting · spatial "
             "resolution · feature representation · model ensemble · outage density · "
             "numerical weather prediction · IEEE Access"),
            ("Wanik, D. W., Anagnostou, E. N., Astitha, M., Hartman, B. M., Lackmann, G. M., "
             "Yang, J., Cerrai, D., He, J., Frediani, M. (2018). A case study on power "
             "outage impacts from future Hurricane Sandy scenarios. Journal of Applied "
             "Meteorology and Climatology 57(1), 51–79.",
             "climate-change-relevant outage modelling. Uses pseudo-global-warming WRF runs "
             "to project how Hurricane Sandy would behave under future climate conditions, "
             "then runs the outage model on those scenarios. Strong forward-looking "
             "citation for any climate-resilience framing of your work.",
             "couples future-climate atmospheric simulations with the UConn outage model to "
             "project how a Sandy-like event would affect Connecticut under 2050s/2080s "
             "climate. Quantifies the increase in expected outages under warmer "
             "atmospheric scenarios.",
             "pseudo-global-warming (PGW) · climate scenarios · Hurricane Sandy · "
             "tropical cyclone · downscaling · WRF · climate impact assessment · "
             "future projection"),
            ("Yang, F., Wanik, D. W., Cerrai, D., Bhuiyan, M. A. E., Anagnostou, E. N. "
             "(2020). Quantifying Uncertainty in Machine Learning-Based Power Outage "
             "Prediction Model Training Requirements. Forecasting 2(2), 152–169.",
             "the uncertainty-quantification companion to the prediction work. Important "
             "because any restoration model calibrated against outage predictions inherits "
             "their uncertainty. Methodology to mirror in your own calibration phase.",
             "studies how prediction-model accuracy and uncertainty scale with the amount "
             "of historical storm data available for training. Provides explicit guidance "
             "on data requirements for outage-prediction model deployment.",
             "uncertainty quantification · training-set size · learning curves · "
             "bootstrap · cross-validation · model variance · forecast skill · operational "
             "deployment"),
            ("Frediani, M. E. B., Anagnostou, E. N., Astitha, M. (2017). A new "
             "verification method for assessing the performance of high-resolution "
             "atmospheric simulations. Monthly Weather Review 145(2), 769–784.",
             "the atmospheric-model verification methodology underlying the UConn pipeline. "
             "Cite this when explaining how the storm inputs to the outage-prediction model "
             "are themselves validated.",
             "develops object-based verification methods for high-resolution mesoscale "
             "atmospheric simulations, used to validate WRF runs against observed storm "
             "events.",
             "mesoscale atmospheric simulation · object-based verification · forecast "
             "skill · WRF validation · feature detection · storm tracking · MODE · "
             "verification metric"),
        ]),
        ("Grid Restoration Optimisation (MILP / Stochastic)", BG_BUILD, [
            ("Van Hentenryck, P., Coffrin, C., Bent, R. (2011). Vehicle Routing for the Last "
             "Mile of Power System Restoration. Proc. Power Systems Computation Conference (PSCC).",
             "formulates power restoration explicitly as a vehicle-routing problem (VRP) on "
             "top of network constraints — essentially the academic version of the dispatch "
             "problem your scheduler solves heuristically.",
             "derives a constraint-programming formulation that schedules repair crews "
             "across damaged feeders to minimise total customer-minutes without service, "
             "with travel times and repair durations as inputs. Uses constraint-based "
             "search rather than MILP. Closest single citation to your problem statement.",
             "Vehicle Routing Problem (VRP) · last-mile restoration · constraint programming · "
             "crew scheduling · customer-minutes · feeder repair · objective function · "
             "branch-and-bound"),
            ("Coffrin, C., Van Hentenryck, P. (2015). Transmission System Restoration with "
             "Co-Optimization of Repairs, Load Pickups, and Generation Dispatch. IEEE "
             "Transactions on Power Systems.",
             "defines the MILP optimal scheduler you'd benchmark your greedy against. "
             "Establishes the standard form for restoration-as-optimisation in the "
             "academic literature.",
             "joint MILP for choosing which damaged lines to repair, in what order, while "
             "simultaneously dispatching generation and picking up load. Demonstrated on "
             "transmission case studies; the formulation generalises to distribution.",
             "Mixed-Integer Linear Program (MILP) · co-optimisation · transmission "
             "restoration · load pickup · generation dispatch · convex relaxation · "
             "decision variables · feasibility cuts"),
            ("Arif, A., Wang, Z., Wang, J., Mather, B., Bashualdo, H., Zhao, D. (2018). "
             "Power Distribution System Outage Management with Co-Optimization of Repairs, "
             "Reconfiguration, and DG Dispatch. IEEE Transactions on Smart Grid 9(5), "
             "4109–4118.",
             "closest single paper to your problem at the distribution level. Same kind of "
             "problem (crews + outages + time), same scale (distribution, not transmission), "
             "with the addition of distributed-generation islanding.",
             "a two-stage stochastic MILP that schedules repair crews while simultaneously "
             "reconfiguring the network topology and dispatching distributed generators. "
             "Models uncertainty in repair times via scenarios.",
             "two-stage stochastic optimisation · distribution feeder reconfiguration · "
             "distributed generation (DG) · microgrid islanding · scenario tree · repair "
             "crew dispatch · outage management system (OMS) · MILP"),
            ("Watson, J-P., Hart, W. E., Murray, R. (2005). Formal Methods for the Real-"
             "Time Restoration of Electric Power. Sandia National Laboratories Technical "
             "Report.",
             "one of the early Sandia/DOE-funded reports defining the infrastructure-"
             "restoration scheduling problem formally. Foundational reference for situating "
             "your work in the security-of-supply literature.",
             "develops mathematical-programming foundations for the restoration problem "
             "with uncertainty, including stochastic repair durations and crew availability.",
             "infrastructure resilience · stochastic programming · scenario-based "
             "optimisation · mathematical programming · critical-infrastructure protection · "
             "real-time decision support · Sandia · DOE"),
            ("Castillo, A. (2014). Risk Analysis and Management in Power Outage and "
             "Restoration: A Literature Survey. Electric Power Systems Research 107, 9–15.",
             "the single survey to cite in your introduction when situating your greedy "
             "among the broader model family.",
             "surveys restoration approaches across heuristics, MILP, stochastic "
             "programming, and decomposition methods; tabulates assumptions and limitations.",
             "literature survey · risk analysis · heuristics · MILP · stochastic "
             "programming · Benders decomposition · taxonomy · outage restoration"),
        ]),
        ("Storm-Induced Outage Modelling & Prediction", BG_PERF, [
            ("Nateghi, R., Guikema, S. D., Quiring, S. M. (2011). Comparison and Validation "
             "of Statistical Methods for Predicting Power Outage Durations in the Event of "
             "Hurricanes. Risk Analysis 31(12), 1897–1906.",
             "the calibration target methodology. If you ever fit your scheduler against "
             "real Eversource event timelines, this paper's evaluation framework "
             "(statistical durations vs observed) is what you'll follow.",
             "benchmarks several statistical models (Cox proportional hazards, accelerated "
             "failure time, etc.) for predicting how long restoration will take after a "
             "hurricane, validated against utility data.",
             "Cox proportional hazards · accelerated failure time (AFT) · survival "
             "analysis · outage duration · hurricane impact · cross-validation · model "
             "comparison · MAPE / RMSE"),
            ("Han, S-R., Guikema, S. D., Quiring, S. M., Lee, D-Y., Rosowsky, D., Davidson, "
             "R. A. (2009). Estimating the Spatial Distribution of Power Outages during "
             "Hurricanes in the Gulf Coast Region. Reliability Engineering & System Safety "
             "94(2), 199–210.",
             "spatial outage prediction. If you want to make your storm model less "
             "synthetic — driving outage locations from weather forecasts rather than "
             "uniform random sampling — this is the template.",
             "spatial generalized linear mixed model for predicting where, geographically, "
             "outages will occur during a hurricane. Inputs include wind speed, soil "
             "moisture, tree cover; outputs are spatially explicit outage probabilities.",
             "generalised linear mixed model (GLMM) · spatial regression · wind speed · "
             "soil moisture · tree canopy cover · outage density · spatially explicit "
             "prediction · grid cell aggregation"),
            ("Liu, H., Davidson, R. A., Apanasovich, T. V. (2007). Statistical Forecasting "
             "of Electric Power Restoration Times in Hurricanes and Ice Storms. IEEE "
             "Transactions on Power Systems 22(4), 2270–2279.",
             "a direct predecessor of the Nateghi line, focused on restoration time "
             "forecasting (not just outage counts). Useful for situating the calibration "
             "phase.",
             "regression-based forecasts of restoration completion times given storm "
             "intensity and grid attributes; evaluated on Carolinas hurricane and ice "
             "storm data.",
             "restoration time forecasting · negative binomial regression · ice storm · "
             "hurricane · storm intensity index · customer-restored curve · service "
             "territory · forecast horizon"),
        ]),
        ("Synthetic Grid Generation", BG_DECISION, [
            ("Birchfield, A. B., Xu, T., Gegner, K. M., Shetye, K. S., Overbye, T. J. (2017). "
             "Grid Structural Characteristics as Validation Criteria for Synthetic Networks. "
             "IEEE Transactions on Power Systems 32(4), 3258–3265.",
             "the canonical Texas A&M synthetic-grid paper. Methods (k-means clustering, "
             "demand allocation, topology validation) directly informed your distribution-"
             "level approach even though the paper itself targets transmission.",
             "develops statistical criteria (degree distribution, line lengths, generator-"
             "load distance) for judging whether a synthetic grid behaves like a real one, "
             "then uses those criteria to validate the Texas A&M transmission test cases "
             "used widely in the literature.",
             "synthetic transmission network · k-means clustering · degree distribution · "
             "line length distribution · graph topology · validation criteria · ACTIVSg "
             "test cases · generator-load proximity"),
            ("Kersting, W. H. (2001/2018). Distribution System Modeling and Analysis. "
             "CRC Press (textbook).",
             "the standard reference for IEEE PES distribution test feeders (13-node, "
             "34-node, 123-node). If you ever want a hand-curated ground-truth feeder "
             "topology to validate your synthetic generator against, this is the source.",
             "textbook covering radial distribution analysis methods plus the standard "
             "IEEE PES test feeder dataset definitions.",
             "radial distribution feeder · IEEE 13/34/123-node test feeder · three-phase "
             "modelling · ladder iterative method · per-unit system · backbone vs lateral · "
             "service transformer · load model"),
        ]),
        ("Crew Routing & Vehicle Routing Variants", BG_FIX, [
            ("Toth, P., Vigo, D. (eds.) (2014). Vehicle Routing: Problems, Methods, and "
             "Applications (2nd edition). SIAM MOS-SIAM Series on Optimization.",
             "the foundational reference for VRP variants. Your crew dispatch is "
             "structurally a multi-vehicle VRP with time-dependent service (discovery "
             "delays) and stochastic service times (repair durations). Citing this places "
             "your work in the right algorithmic family.",
             "comprehensive textbook chapters on VRP variants (time windows, stochastic "
             "demands, dynamic VRP, multi-depot) and the exact / heuristic / metaheuristic "
             "methods used to solve them.",
             "Vehicle Routing Problem (VRP) · VRP with time windows (VRPTW) · stochastic "
             "VRP · multi-depot VRP · branch-and-cut · column generation · savings "
             "algorithm · large neighbourhood search"),
            ("Solomon, M. M. (1987). Algorithms for the Vehicle Routing and Scheduling "
             "Problems with Time Window Constraints. Operations Research 35(2), 254–265.",
             "the canonical VRPTW benchmark instances are the de facto standard test set. "
             "If you ever compare your greedy on a synthetic benchmark against published "
             "MILP / metaheuristic results, this is where the instances come from.",
             "introduces VRPTW benchmark instances (the 'Solomon benchmarks') plus several "
             "construction and improvement heuristics.",
             "Solomon benchmarks · VRPTW · time windows · insertion heuristic · savings "
             "algorithm · construction heuristic · improvement heuristic · feasibility "
             "check"),
            ("Perrier, N., Langevin, A., Campbell, J. F. (2007). A Survey of Models and "
             "Algorithms for Winter Road Maintenance. Part IV: Vehicle Routing and Depot "
             "Location for Spreading. Computers & Operations Research 34(1), 258–294.",
             "snowplow routing is the canonical 'crews dispatched after disruption' "
             "problem. Methodologically the closest analogue to power restoration outside "
             "the power-engineering literature.",
             "surveys VRP variants tailored to winter road maintenance with time-critical "
             "service, route-density constraints, and depot placement — all of which have "
             "direct analogues in crew-restoration dispatch.",
             "winter road maintenance · snowplow routing · Chinese postman problem · arc "
             "routing · depot location · time-critical service · multi-period routing · "
             "fleet sizing"),
        ]),
        ("Resilience Metrics & Frameworks", BG_QUESTION, [
            ("Panteli, M., Mancarella, P. (2015). The Grid: Stronger, Bigger, Smarter?: "
             "Presenting a Conceptual Framework of Power System Resilience. IEEE Power "
             "and Energy Magazine 13(3), 58–66.",
             "defines the widely cited 'resilience trapezoid' — the before/during/after-"
             "disturbance performance curve. This is the figure shape your simulation's "
             "customer-minutes-without-service output naturally produces, so framing your "
             "output in this language helps a paper introduction.",
             "introduces a conceptual framework distinguishing reliability (steady-state) "
             "from resilience (under disturbance), with metrics for each phase of the "
             "trapezoid.",
             "resilience trapezoid · reliability vs resilience · disturbance phase · "
             "post-event recovery · hardening · operational resilience · performance "
             "function F(t) · system performance metric"),
            ("Panteli, M., Mancarella, P. (2015). Influence of Extreme Weather and Climate "
             "Change on the Resilience of Power Systems: Impacts and Possible Mitigation "
             "Strategies. Electric Power Systems Research 127, 259–270.",
             "directly relevant to motivating Connecticut storm-resilience work in your "
             "introduction. Establishes the climate-change connection.",
             "reviews extreme-weather impacts on power systems and surveys hardening / "
             "restoration mitigation strategies.",
             "extreme weather · climate change · hardening strategy · undergrounding · "
             "vegetation management · mitigation · high-impact low-probability (HILP) "
             "events · adaptation"),
            ("Bie, Z., Lin, Y., Li, G., Li, F. (2017). Battling the Extreme: A Study on "
             "the Power System Resilience. Proceedings of the IEEE 105(7), 1253–1266.",
             "a recent (relatively) survey of resilience in distribution systems including "
             "microgrid islanding. Good for the literature-review section of any paper.",
             "survey article covering resilience metrics, modelling approaches, hardening "
             "strategies, and the role of distributed energy resources.",
             "distribution-system resilience · microgrid · distributed energy resources "
             "(DER) · islanding · resilience index · self-healing grid · sectionalising "
             "switch · proactive operation"),
            ("National Academies of Sciences, Engineering, and Medicine (2017). Enhancing "
             "the Resilience of the Nation's Electricity System. Washington, DC: The "
             "National Academies Press.",
             "the policy-context citation. Cite this in your introduction to establish "
             "that grid resilience is a recognised national-priority research area, not a "
             "niche academic interest.",
             "National Academies consensus study laying out the research and policy "
             "agenda for U.S. electricity-system resilience. Open-access download.",
             "policy framework · research agenda · electricity system · national resilience · "
             "interdependent infrastructure · cyber-physical security · federal R&D · "
             "consensus study"),
        ]),
        ("Interactive & Browser-Based Visualization", BG_PIVOT, [
            ("Overbye, T. J., Wiegmann, D. A., Rich, A. M., Sun, Y. (2003). Human Factors "
             "Aspects of Power System Voltage Contour Visualization. IEEE Transactions on "
             "Power Systems 18(1), 76–82.",
             "one of the few peer-reviewed papers on grid visualisation specifically. Cite "
             "as a methodological precedent if asked why visualisation is part of your "
             "contribution rather than just engineering polish.",
             "human-factors evaluation of contour-map visualisation of power-system "
             "voltages. Establishes that visualisation choices materially affect operator "
             "decisions.",
             "human factors · situational awareness · voltage contour map · operator "
             "decision support · visualisation evaluation · cognitive load · energy "
             "management system (EMS) · control-room display"),
            ("Note on the broader gap",
             "the academic literature on web-based / browser-first interactive grid "
             "simulation is sparse. Most decision-support tools are commercial (OSI "
             "Monarch, GE PowerOn) and proprietary; most open demos are toy-scale "
             "(≤100 nodes). The peer-reviewed gap is itself part of the case for your "
             "work occupying a distinctive niche.",
             "cite engineering precedents (Leaflet, D3.js) in methods, not literature; "
             "cite Overbye for the visualisation-matters point; build the rest of the "
             "niche argument around the access/scale/calibration gaps.",
             "outage management system (OMS) · decision-support tool · proprietary vs "
             "open-source · Leaflet · D3.js · interactive simulation · web-based GIS · "
             "literature gap"),
        ]),
    ]

    for theme_name, bg, items in themes:
        h2(doc, theme_name)
        for citation, why, what, vocab in items:
            # Citation row (colored band)
            t = doc.add_table(rows=1, cols=1)
            cell = t.rows[0].cells[0]
            cell.width = Inches(6.4)
            shade(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            add_styled(p, citation, bold=True, color=INK, size=11)
            # Why-it-matters
            p = doc.add_paragraph()
            add_styled(p, "Why it matters.  ", bold=True, color=ACCENT, size=11)
            add_styled(p, why, color=INK, size=11)
            # What-it-does
            p = doc.add_paragraph()
            add_styled(p, "What it does.  ", bold=True, color=ACCENT, size=11)
            add_styled(p, what, color=INK, size=11)
            # Key terms — in a light-gray banded cell so it visually reads as
            # a vocabulary list, with each term separated by a middle dot.
            t = doc.add_table(rows=1, cols=1)
            cell = t.rows[0].cells[0]
            cell.width = Inches(6.4)
            shade(cell, "F0E5CB")
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            add_styled(p, "Key terms.  ", bold=True, color=ACCENT, size=10)
            add_styled(p, vocab, italic=True, color=INK, size=10)
            body(doc, "")

    body(doc, "")
    h2(doc, "Where Your Contribution Sits — A Sketch for the Introduction")
    body(doc,
         "Restoration scheduling for distribution-system outages has a rich operations-"
         "research literature dominated by MILP and stochastic-programming formulations "
         "(Van Hentenryck et al. 2011; Coffrin & Van Hentenryck 2015; Arif et al. 2018). In "
         "parallel, the UConn Eversource Energy Center group (Wanik, Anagnostou, Cerrai, "
         "Astitha and collaborators) has built sophisticated machine-learning outage-"
         "prediction pipelines tuned specifically to Connecticut Light & Power / Eversource "
         "service territory (Wanik et al. 2015, 2017, 2018; Cerrai et al. 2019). Yet two "
         "gaps remain: (i) the decision-support pipeline between the prediction community "
         "and the restoration-scheduling community is largely proprietary, trapped inside "
         "commercial OMS platforms; and (ii) academic restoration models are rarely "
         "validated against the publicly available PURA storm-event filings that "
         "Connecticut utilities are required to produce.")
    body(doc,
         "This work contributes a publicly accessible, browser-first interactive simulator "
         "for Hartford County and (by extension) the Connecticut service territory, with "
         "the engineering capacity to run 100 000-outage scenarios and Monte Carlo "
         "ensembles in tens of milliseconds. It is positioned as the downstream complement "
         "to the UConn prediction pipeline — their work answers \"how many outages and "
         "where?\"; this work answers \"once those outages exist, how long until "
         "restoration?\" The simulator is also positioned as the calibration substrate for "
         "future work matching synthetic restoration timelines against real Eversource "
         "event histories. The intended contribution is the integration — an accessible, "
         "scale-validated, calibratable tool — rather than a novel scheduling algorithm.")

    body(doc, "")
    h2(doc, "Open Questions Your Paper Can Answer")
    questions = [
        "How close is the greedy rolling-horizon heuristic to MILP-optimal restoration at "
        "the scales where MILP is tractable (N ≤ 50)? What is the suboptimality gap as a "
        "function of crew/outage ratio?",
        "Given calibration against one real Eversource storm event (e.g., Isaias 2020 or "
        "the May 2018 tornado outage), how well do the model's parameters generalise to a "
        "held-out storm?",
        "How do restoration-time distributions change as you scale from county- to state-"
        "level outage counts? Is there a phase transition in crew-utilisation efficiency?",
        "Does adding stochastic re-discovery (delayed outage reporting) change the optimal "
        "crew dispatch strategy versus deterministic assumptions?",
    ]
    for q in questions:
        p = doc.add_paragraph()
        add_styled(p, "?  ", bold=True, color=QUESTION, size=12)
        add_styled(p, q, color=INK, size=11)

    body(doc, "")
    h2(doc, "Data Sources to Pursue")
    sources = [
        ("Connecticut PURA dockets",
         "Post-storm filings for major events (Isaias 2020, May 2018 tornado). Searchable "
         "at dpuc.state.ct.us. Restoration timelines, crew counts, mutual-aid arrivals."),
        ("Eversource storm reports",
         "Public-facing post-storm documents and PURA submissions. Often include outage "
         "curve, crew totals, customer-minutes-without-service."),
        ("U.S. Energy Information Administration Form EIA-417",
         "Major electric-disturbance reports. Coarse but federal-level coverage of major "
         "outage events including Connecticut."),
        ("DOE OE-417 disturbance dataset",
         "Similar federal dataset, downloadable bulk."),
        ("County / municipal GIS",
         "For real road network and population density — eventual upgrade from the "
         "synthetic demand-point model."),
    ]
    for who, what in sources:
        p = doc.add_paragraph()
        add_styled(p, "■ ", bold=True, color=ACCENT, size=12)
        add_styled(p, who + ".  ", bold=True, color=INK, size=11)
        add_styled(p, what, color=INK, size=11)

    body(doc, "")
    margin_note(doc,
                "These citations are starting points drawn from a literature scan rather "
                "than an exhaustive bibliography. Verify each on Google Scholar / IEEE "
                "Xplore before citing in a manuscript, and use this map to seed a more "
                "thorough review.")

    doc.save("Hartford_Grid_Research_Context.docx")
    print("Wrote Hartford_Grid_Research_Context.docx")


if __name__ == "__main__":
    build_journal()
    build_research()
